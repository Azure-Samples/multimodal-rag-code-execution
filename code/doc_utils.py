#   ________  .__         ___.            .__       __________ .__                    __        __________         .__    __       
#  /  _____/  |  |   ____ \_ |__  _____   |  |      \______   \|  |  _____     ____  |  | __    \______   \  ____  |  | _/  |_     
# /   \  ___  |  |  /  _ \ | __ \ \__  \  |  |       |    |  _/|  |  \__  \  _/ ___\ |  |/ /     |    |  _/_/ __ \ |  | \   __\    
# \    \_\  \_|  | (  <_> )| \_\ \ / __ \_|  |__     |    |   \|  |__ / __ \_\  \___ |    <      |    |   \\  ___/ |  |__|  |      
#  \______  / |____/\____/ |___  /(____  /|____/     |______  /|____/(____  / \___  >|__|_ \     |______  / \___  >|____/|__|      
#         \/             \/      \/                   \/            \/      \/      \/            \/      \/                 
#                                                        .__                                .__ .__            __            
# _______   ____    ______  ____  _____  _______   ____  |  |__        ____   ____  ______  |__||  |    ____ _/  |_          
# \_  __ \_/ __ \  /  ___/_/ __ \ \__  \ \_  __ \_/ ___\ |  |  \     _/ ___\ /  _ \ \____ \ |  ||  |   /  _ \\   __\         
#  |  | \/\  ___/  \___ \ \  ___/  / __ \_|  | \/\  \___ |   Y  \    \  \___(  <_> )|  |_> >|  ||  |__(  <_> )|  |           
#  |__|    \___  >/____  > \___  >(____  /|__|    \___  >|___|  /     \___  >\____/ |   __/ |__||____/ \____/ |__|           
#              \/      \/      \/      \/             \/      \/          \/        |__|                                     

#  /\_/\
# ( o.o )
#  > ^ <

# /~\
# C oo
#  _( ^)
# /   ~\

# AML CHEATSHEET
# https://azure.github.io/azureml-cheatsheets/docs/cheatsheets/python/v1/compute-targets
import fitz  # PyMuPDF
import os
import glob
import traceback
from dotenv import load_dotenv
import shutil
import copy
import json
import uuid
import hashlib
import pandas as pd
from utils.bcolors import bcolors as bc 
import tiktoken
from multiprocessing.dummy import Pool as ThreadPool
import math
import itertools 
import time
import numpy as np
import time
import random
from PIL import Image

from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult
from azure.ai.documentintelligence.models import ContentFormat

load_dotenv()

from env_vars import *
from utils.logc import logc, get_current_time
from utils.http_helpers import *
from utils.cogsearch_rest import *
from utils.sc_sync import *
from utils.file_utils import (
    is_file_or_url, 
    save_to_pickle, 
    write_to_file, 
    read_asset_file, 
    replace_extension, 
    download_file, 
    check_replace_extension, 
    generate_file_sas_full_link)
from utils.text_utils import (
    extract_code, get_token_count, limit_token_count, recover_json, extract_markdown, extract_markdown_table, 
    extract_markdown_table_as_df, extract_extracted_text, remove_extracted_text, remove_code, extract_mermaid, 
    extract_chunk_number, clean_up_text)
from utils.openai_utils import (
    gpt4_models, 
    get_chat_completion, 
    get_chat_completion_with_json, 
    get_embeddings, 
    ask_LLM, 
    ask_LLM_with_JSON, 
    call_gpt4v)
from utils.code_exec_utils import (
    execute_python_code_block, 
    try_code_interpreter_for_tables_using_python_exec, 
    try_code_interpreter_for_tables_using_taskweaver)
from utils.assistants_utils import try_code_interpreter_for_tables_using_assistants_api

# Begin main

try:
    from docx import Document
except:
    print("WARNING: docx module not found. Please install python-docx")

try:
    from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
    from llama_index.core import SimpleDirectoryReader
    from llama_index.core import Settings
    from llama_index.core.node_parser import (
        SentenceSplitter,
        SemanticSplitterNodeParser,
    )

except:
    print("WARNING: one of the llama_index module(s) not found. Please install llama_index modules.")

pool = ThreadPool(20)


def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))



if os.path.exists("./prompts"):
    prompt_dir = "./prompts"
else:
    prompt_dir = "../code/prompts"

section_generation_prompt = read_asset_file(f'{prompt_dir}/section_generation_prompt.txt')[0]
general_prompt_template = read_asset_file(f'{prompt_dir}/general_prompt_template.txt')[0]
specialized_prompt_template = read_asset_file(f'{prompt_dir}/specialized_prompt_template.txt')[0]
numerical_prompt_template = read_asset_file(f'{prompt_dir}/numerical_prompt_template.txt')[0]
table_prompt_template = read_asset_file(f'{prompt_dir}/table_prompt_template.txt')[0]
image_prompt_template = read_asset_file(f'{prompt_dir}/image_prompt_template.txt')[0]
rate_answers_prompt_template = read_asset_file(f'{prompt_dir}/rate_answers_prompt_template.txt')[0]
code_harvesting_from_text = read_asset_file(f'{prompt_dir}/code_harvesting_from_text.txt')[0]
markdown_extract_header_and_summarize_prompt = read_asset_file(f'{prompt_dir}/markdown_extract_header_and_summarize_prompt.txt')[0]
process_extracted_text_prompt = read_asset_file(f'{prompt_dir}/process_extracted_text_prompt.txt')[0]
chunk_analysis_template = read_asset_file(f'{prompt_dir}/chunk_analysis_template.txt')[0]
extract_text_from_images_prompt = read_asset_file(f'{prompt_dir}/extract_text_from_images_prompt.txt')[0]
vision_system_prompt = read_asset_file(f'{prompt_dir}/vision_system_prompt.txt')[0]
detect_num_of_tables_prompt = read_asset_file(f'{prompt_dir}/detect_num_of_tables_prompt.txt')[0]
detect_num_of_diagrams_prompt = read_asset_file(f'{prompt_dir}/detect_num_of_diagrams_prompt.txt')[0]
image_description_prompt = read_asset_file(f'{prompt_dir}/image_description_prompt.txt')[0]
table_code_description_prompt = read_asset_file(f'{prompt_dir}/table_code_description_prompt.txt')[0]
table_markdown_description_prompt = read_asset_file(f'{prompt_dir}/table_markdown_description_prompt.txt')[0]
table_qa_prompt = read_asset_file(f'{prompt_dir}/table_qa_prompt.txt')[0]
context_extension = read_asset_file(f'{prompt_dir}/context_extension.txt')[0]
optimize_embeddings_prompt = read_asset_file(f'{prompt_dir}/optimize_embeddings_prompt.txt')[0]
document_wide_tags = read_asset_file(f'{prompt_dir}/document_wide_tags.txt')[0]
document_wide_summary = read_asset_file(f'{prompt_dir}/document_wide_summary.txt')[0]
user_query = read_asset_file(f'{prompt_dir}/user_query.txt')[0]
user_query = read_asset_file(f'{prompt_dir}/user_query.txt')[0]
direct_user_query = read_asset_file(f'{prompt_dir}/direct_user_query.txt')[0]
table_info = read_asset_file(f'{prompt_dir}/table_info.txt')[0]
py_files_import = read_asset_file(f'{prompt_dir}/py_files_import.txt')[0]
search_context_extension = read_asset_file(f'{prompt_dir}/search_context_extension.txt')[0]
summaries_context_extension = read_asset_file(f'{prompt_dir}/summaries_context_extension.txt')[0]
search_system_prompt = read_asset_file(f'{prompt_dir}/search_system_prompt.txt')[0]
search_prompt = read_asset_file(f'{prompt_dir}/search_prompt.txt')[0]
full_search_json_output = read_asset_file(f'{prompt_dir}/full_search_json_output.txt')[0]
limited_search_json_output = read_asset_file(f'{prompt_dir}/limited_search_json_output.txt')[0]
computation_is_needed_prompt = read_asset_file(f'{prompt_dir}/computation_is_needed_prompt.txt')[0]
vision_support_prompt = read_asset_file(f'{prompt_dir}/vision_support_prompt.txt')[0]
query_entities_prompt = read_asset_file(f'{prompt_dir}/query_entities_prompt.txt')[0]
search_learnings_template = read_asset_file(f'{prompt_dir}/search_learnings_template.txt')[0]
detect_intent_prompt = read_asset_file(f'{prompt_dir}/detect_intent_prompt.txt')[0]







def generate_uuid_from_string(input_string):
    # Create a SHA-1 hash of the input string
    hash_object = hashlib.sha1(input_string.encode())
    # Use the first 16 bytes of the hash to create a UUID
    return str(uuid.UUID(bytes=hash_object.digest()[:16]))
    


def generate_section(section):
    prompt = section_generation_prompt.format(section=section)
    return ask_LLM(prompt)

def get_solution_path_components(asset_path: str) -> dict:
    stem = os.path.normpath(asset_path)
    components = []

    for i in range(4):
        arr_split = os.path.split(stem)
        components.append(arr_split[-1])
        stem = arr_split[0]

    dd = { "index_name": components[3], "proc_filename": components[2], "type_dir": components[1], "filename": components[0]}

    return dd

def copy_ai_search_index(source_index, target_index):
    sc = IndexSync(staging_index_name = source_index, prod_index_name = target_index)
    sc.prod_cs.create_index()
    output = sc.duplicate_index()
    logc("Copy AI Search Index", f"Source Index: {source_index}. Target Index: {target_index}. Status: Copied {len(output)} items.")
    return len(output)

def update_document_in_index(asset_file, new_doc, index_name):
    index = CogSearchRestAPI(index_name)

    dd = get_solution_path_components(asset_file)

    unique_identifier = f"{dd['index_name']}_{dd['proc_filename']}_{os.path.basename(asset_file)}"
    asset_id = generate_uuid_from_string(unique_identifier)
    new_doc["asset_id"] = asset_id

    doc = index.get_document_by_id(asset_id)
    if doc is None:
        logc("Error retrieving document from index", f"Document {asset_id} not found in index {index_name}")
        return None

    copy_doc = copy.deepcopy(doc)
    copy_doc['vector'] = "<vector>"
    logc("Document found ", json.dumps(copy_doc, indent=4))



    unique_identifier = f"{dd['index_name']}_{dd['proc_filename']}"
    document_id = generate_uuid_from_string(unique_identifier)
    
    if doc['document_id'] != document_id: 
        logc("Error updating document in index", f"Document {asset_id} is not associated with the correct pdf document {document_id}")
        return None
   
    for f in ["@odata.context", "@search.score", "@search.rerankerScore", "@search.captions"]:
        if doc.get(f, None) is not None:
            del doc[f]

    for key in new_doc:
        doc[key] = new_doc[key]

    copy_doc = copy.deepcopy(doc)
    copy_doc['vector'] = "<vector>"
    logc("New document to be uploaded ", json.dumps(copy_doc, indent=4))

    res = index.upload_documents([doc])

    return res

def get_dpi(image):
    try:
        dpi = image.info['dpi']
        # print("DPI", dpi)
    except KeyError:
        dpi = (300, 300)
    return dpi

def polygon_to_bbox(polygon):
    xs = polygon[::2]  # Extract all x coordinates
    ys = polygon[1::2]  # Extract all y coordinates
    left = min(xs)
    top = min(ys)
    right = max(xs)
    bottom = max(ys)
    return (left, top, right, bottom)

def polygons_to_bbox(polygons):
    xs = [point for polygon in polygons for point in polygon[::2]]
    ys = [point for polygon in polygons for point in polygon[1::2]]
    left = min(xs)
    top = min(ys)
    right = max(xs)
    bottom = max(ys)
    return (left, top, right, bottom)

def inches_to_pixels(inches, dpi):
    dpi_x, dpi_y = dpi
    return [int(inches[i] * dpi_x if i % 2 == 0 else inches[i] * dpi_y) for i in range(len(inches))]

def extract_figure(image_path, polygons, target_filename):
    bbox_in_inches = polygons_to_bbox(polygons)

    # Load the image
    image = Image.open(image_path)
    # print(f"Cropped image will be saved under name: {target_filename}")

    # Get DPI from the image
    dpi = get_dpi(image)

    # Convert the bounding box to pixels using the image's DPI
    bbox_in_pixels = inches_to_pixels(bbox_in_inches, dpi)
    # print("bbox_in_pixels", bbox_in_pixels)

    # Crop the image
    cropped_image = image.crop(bbox_in_pixels)
    cropped_image.save(target_filename)  

    return target_filename

def get_excel_sheet_names(file_path):
    # Load the Excel file
    xls = pd.ExcelFile(file_path, engine='openpyxl')

    # Get the list of sheet names
    sheet_names = xls.sheet_names

    return sheet_names

def read_csv_to_dataframe(file_path):
    # Read a CSV file into a DataFrame
    df = pd.read_csv(file_path)

    return {"sheet1": df}

def read_excel_to_dataframes(file_path):
    xls = pd.ExcelFile(file_path, engine='openpyxl')
    dfs = {}

    # Read each sheet into a DataFrame
    for sheet_name in xls.sheet_names:
        dfs[sheet_name] = pd.read_excel(xls, sheet_name, header=None)

    return dfs

def table_df_cleanup_df(df):
    dfc = copy.deepcopy(df)
    dfc = dfc.dropna(axis=0, how='all')
    dfc = dfc.dropna(axis=1, how='all')
    dfc = dfc.replace(r'\n','   //    ', regex=True) 
    dfc = dfc.replace(r'\|','   ///    ', regex=True) 
    return dfc

# def save_pptx_as_pdf(pptx_path, output_directory):
#     comtypes.CoInitialize()

#     # Path to the output PDF file
#     base_name = os.path.splitext(os.path.basename(pptx_path))[0]
#     extension = os.path.splitext(os.path.basename(pptx_path))[1]

#     output_document_path = os.path.join(output_directory, base_name + '.pdf')

#     print("Basename: ", base_name)
#     print("Extension: ", extension)
#     print("Output PDF Path: ", output_document_path)
    
#     try:
#         try:
#             # Initialize PowerPoint
#             powerpoint = comtypes.client.CreateObject("Powerpoint.Application")

#             # Open the PPTX file
#             presentation = powerpoint.Presentations.Open(pptx_path)

#             # Save the presentation as a PDF
#             presentation.SaveAs(output_document_path, FileFormat=32)  # 32 corresponds to the PDF format in PowerPoint
#             description = f"The file has been successfully converted to PDF and saved at {output_document_path}."
        
#         finally:
#             # Close the presentation and PowerPoint
#             presentation.Close()
#             powerpoint.Quit()

#     except Exception as e:
#         description = f"The file could not be converted to PDF.\n{e}"
#         output_document_path = ''


#     # Release COM objects
#     comtypes.CoUninitialize()

#     print("Status: ", description)

#     # Result variable
#     return output_document_path, 



# def save_docx_as_pdf(docx_path, output_directory):

#     comtypes.CoInitialize()
    
#     # Convert the .docx file to PDF using Microsoft Word's COM API
#     word = comtypes.client.CreateObject('Word.Application')
#     word.Visible = False

#     base_name = os.path.splitext(os.path.basename(docx_path))[0]
#     extension = os.path.splitext(os.path.basename(docx_path))[1]

#     output_document_path = os.path.join(output_directory, base_name + '.pdf')

#     print("\n\nDocx Path: ", docx_path)
#     print("Basename: ", base_name)
#     print("Extension: ", extension)
#     print("Output PDF Path: ", output_document_path)

#     try:
#         try:
#             doc = word.Documents.Open(os.path.abspath(docx_path))
#             print(f"Word Document successfully opened. {os.path.abspath(docx_path)}")
#             doc.SaveAs(os.path.abspath(output_document_path), FileFormat=17)  
#             # FileFormat=17 is for PDFs
#             print(f"PDF Document successfully saved. {os.path.abspath(output_document_path)}")
#             description = f"The file has been successfully converted to PDF and saved at {output_document_path}."      
#         finally:
#             doc.Close()
#             word.Quit()
                
#     except Exception as e:
#         description = f"The file could not be converted to PDF.\n{e}"
#         output_document_path = ''

#     # Release COM objects
#     comtypes.CoUninitialize()

#     print("Status: ", description)
#     return output_document_path





def chunk_markdown_table_with_overlap(md_table, cols = None, n_tokens = 512, overlap = 128):

    mds = md_table.split('\n')

    if cols is not None:
        header = '|   ' + '   |   '.join(cols) + '   |\n'
    else:
        header = mds[0] + '\n'

    chunks = []
    chunk = header

    for i, r in enumerate(mds[1:]):
        chunk += r + '\n'

        ## Check if the chunk is over n_tokens
        if get_token_count(chunk) > n_tokens:
            ## Add Overlap
            try:
                for j, ovr in enumerate(mds[i + 2:]):
                    chunk += ovr + '\n'
                    if get_token_count(chunk) > n_tokens + overlap:
                        break
            except Exception as e:
                print(e)
            
            chunks.append(chunk)        

            # print(f"Chunk {len(chunks)}: {get_token_count(chunk)}")
            chunk = header  + mds[1] + '\n'

    return chunks, header

def chunk_markdown_table(md_table, model_info, n_tokens = 512, overlap = 128):
    prompt = markdown_extract_header_and_summarize_prompt.format(table=md_table.split('\n')[:100])
    output = ask_LLM_with_JSON(prompt, model_info = model_info)
    try:
        outd = recover_json(output)
        cols = outd['columns'].split(',')
        summary = outd['summary_of_the_table']
    except:
        try:
            output = ask_LLM_with_JSON(prompt, model_info = model_info)
            outd = recover_json(output)
            cols = outd['columns'].split(',')
            summary = outd['summary_of_the_table']
        except:
            logc(f"Could not recover with malformed JSON {output}")
            return [], '', ''

    chunks, header = chunk_markdown_table_with_overlap(md_table, cols, n_tokens = n_tokens, overlap = overlap)
    print("Chunks:", len(chunks))

    return chunks, header, summary

def chunk_df_as_markdown_table(df, model_info, n_tokens = 512, overlap = 128):
    df_clean = table_df_cleanup_df(df)
    md_table = df_clean.to_markdown()
    chunks, header, summary = chunk_markdown_table(md_table, model_info, n_tokens = n_tokens, overlap = overlap)
    return chunks, header, summary

def write_df_to_files(df, tables_folder, table_count):
    table_path_md = os.path.join(tables_folder, f'chunk_{table_count}_table_0.md')
    table_path = os.path.join(tables_folder, f'chunk_{table_count}_table_0.txt')
    table_path_py = os.path.join(tables_folder, f'chunk_{table_count}_table_0.py')
    write_to_file(df.to_markdown(), table_path_md, 'w')
    write_to_file(df.to_markdown(), table_path, 'w')
    py_script = f"df_{generate_uuid_from_string(str(df.to_dict())).replace('-', '_')} = pd.DataFrame.from_dict({df.to_dict()})"
    write_to_file(py_script, table_path_py, 'w')

    return table_path, table_path_md, table_path_py

def extract_xlsx_using_openpyxl(ingestion_pipeline_dict):
    doc_path = ingestion_pipeline_dict['document_path'] 
    images_folder = ingestion_pipeline_dict['images_directory'] 
    tables_folder = ingestion_pipeline_dict['tables_directory']
    full_text_file = ingestion_pipeline_dict['full_text_file']


    tables = []
    tables_dfs = []
    tables_md = []

    table_count = 0
    
    if ingestion_pipeline_dict['extension'] == '.csv':
        dataframes= read_csv_to_dataframe(doc_path)
    elif ingestion_pipeline_dict['extension'] == '.xlsx':
        dataframes = read_excel_to_dataframes(doc_path)
    else:
        dataframes = {}
        logc("WARNING", f"Could not read the file {doc_path} as it is not a CSV or XLSX file.")

    full_text = ''

    for sheet_name, df in dataframes.items():
        # chunks, header, summary = chunk_df_as_markdown_table(df, model_info, n_tokens = n_tokens, overlap = overlap)
        df_clean = table_df_cleanup_df(df)
        table_path, table_path_md, table_path_py = write_df_to_files(df_clean, tables_folder, table_count)
        tables_dfs.append(table_path_py)
        tables_md.append(table_path_md)
        tables.append(table_path)
        table_count += 1

        full_text += read_asset_file(table_path_md)[0] +'\n\n\n'

    write_to_file(full_text, full_text_file, 'w')

    ingestion_pipeline_dict['tables_py'] = tables_dfs
    ingestion_pipeline_dict['tables_md'] = tables_md
    ingestion_pipeline_dict['table_files'] = tables

    return ingestion_pipeline_dict

def sentence_chunk_text_file(file_path, chunk_size=512, chunk_overlap = 80, model='gpt-4', verbose = False):

    text_splitter = SentenceSplitter(
        separator=" ",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        paragraph_separator="\n\n\n",
        secondary_chunking_regex="[^,.;。]+[,.;。]?",
        tokenizer=tiktoken.encoding_for_model(model).encode
    )

    documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    nodes = text_splitter.get_nodes_from_documents(documents)

    return [n.get_content() for n in nodes]

def semantic_chunk_text_file(file_path, buffer_size=1, breakpoint_percentile_threshold=95, verbose = False):

    embed_model = AzureOpenAIEmbedding(
        model=AZURE_OPENAI_EMBEDDING_MODEL,
        deployment_name=AZURE_OPENAI_EMBEDDING_MODEL,
        api_key=AZURE_OPENAI_EMBEDDING_MODEL_RESOURCE_KEY,
        azure_endpoint=AZURE_OPENAI_EMBEDDING_API_BASE,
        api_version=AZURE_OPENAI_EMBEDDING_MODEL_API_VERSION,
    )

    documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    splitter = SemanticSplitterNodeParser(buffer_size=buffer_size, breakpoint_percentile_threshold=breakpoint_percentile_threshold, embed_model=embed_model)
    nodes = splitter.get_nodes_from_documents(documents)

    if verbose: logc(f"Semantic Chunking Step", f"Number of nodes: {len(nodes)} from file {file_path}")

    return [n.get_content() for n in nodes]

def create_text_doc_chunks_with_sentence_chunking(ingestion_pipeline_dict):

    chunks = ingestion_pipeline_dict.get('chunks', [])
    chunk_index = len(chunks)

    text_directory = ingestion_pipeline_dict['text_directory']
    n_tokens = ingestion_pipeline_dict['chunk_size']
    overlap = ingestion_pipeline_dict['chunk_overlap']

    text_files = []

    try:
        text_chunks = sentence_chunk_text_file(ingestion_pipeline_dict['full_text_file'],
                                               chunk_size=n_tokens,
                                               chunk_overlap=overlap
                                            )
    except:
        text_chunks = []
        logc("Warning", f"Could not perform Sentence Chunking of file {ingestion_pipeline_dict['full_text_file']}")

    
    for index, tc in enumerate(text_chunks):
        text_filename = os.path.join(text_directory, f"chunk_{index}.txt")
        write_to_file(tc, text_filename, 'w')
        chunks.append({
            'chunk_number':chunk_index+1, 
            'text_file': text_filename,
            'full_chunk_text':'',
            'images': [],
            'tables': [],
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'image_mm': [],
            'image_text': [],
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],
            'type': 'text'
        })

        text_files.append(text_filename)
        chunk_index += 1

    ingestion_pipeline_dict['chunks'] = chunks

    ingestion_pipeline_dict['text_files'] = text_files

    return ingestion_pipeline_dict

def create_image_doc_chunks(ingestion_pipeline_dict):

    chunks = ingestion_pipeline_dict.get('chunks', [])
    chunk_index = len(chunks)

    images_directory = ingestion_pipeline_dict['images_directory'] 
    images = ingestion_pipeline_dict['image_files']
    image_text_files = []


    for index, tc in enumerate(images):   
        chunks.append({
            'chunk_number':chunk_index+1, 
            'text_file': '',
            'full_chunk_text':'',
            'chunk_image_path': tc,
            'images': [tc],
            'tables': [],
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'image_mm': [],
            'image_text': [],            
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],        
            'post_process_image_with_context': False, 
            'type': 'image'
        })

        chunk_index += 1

    ingestion_pipeline_dict['chunks'] = chunks

    return ingestion_pipeline_dict

def create_table_doc_chunks_markdown(ingestion_pipeline_dict):

    chunks = ingestion_pipeline_dict.get('chunks', [])
    chunk_index = len(chunks)

    tables_directory = ingestion_pipeline_dict['tables_directory']
    tables_dfs = ingestion_pipeline_dict['tables_py'] 
    tables_md = ingestion_pipeline_dict['tables_md'] 
    tables = ingestion_pipeline_dict['table_files'] 

    model_info = ingestion_pipeline_dict['models'][0]
    n_tokens = ingestion_pipeline_dict['chunk_size']
    overlap = ingestion_pipeline_dict['chunk_overlap']

    text_files = []
    table_text_files = []


    for index, tc in enumerate(tables):
        md_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.md")
        py_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.py")
        table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.txt")

        md_table = read_asset_file(md_filename)[0]
        table_text = read_asset_file(table_text_filename)[0]

        if get_token_count(md_table) > 2 * n_tokens:
            ## First Backup the original table
            md_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.md.backup")
            table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.txt.backup")
            write_to_file(md_table, md_filename, 'w')
            write_to_file(table_text, table_text_filename, 'w')
            
            ## Table too big needs to be broken into chunks
            md_chunks, header, summary = chunk_markdown_table(md_table, model_info, n_tokens = n_tokens, overlap = overlap)

            for md_index, md_chunk in enumerate(md_chunks):
                md_filename = os.path.join(tables_directory, f"chunk_{index}_table_{md_index}.md")
                table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_{md_index}.txt")
                table_text_files.append(table_text_filename)

                text_contents = f"Table Summary:\n{summary}\n\nTable:\n{md_chunk}"
                write_to_file(md_chunk, md_filename, 'w')
                write_to_file(text_contents, table_text_filename, 'w')

                chunks.append({
                    'chunk_number':chunk_index+1, 
                    'text_file': table_text_filename,
                    'full_chunk_text':'',
                    'images': [],
                    'tables': [],
                    'image_py': [],
                    'image_codeblock': [],
                    'image_markdown': [],
                    'image_mm': [],
                    'image_text': [],
                    'table_text': [table_text_filename],
                    'table_py': [py_filename],
                    'table_codeblock': [],
                    'table_markdown': [md_filename],    
                    'type': 'table'    
                })

                chunk_index += 1
        else:        
            table_text_files.append(table_text_filename)

            chunks.append({
                'chunk_number':chunk_index+1, 
                'text_file': table_text_filename,
                'full_chunk_text':'',
                'images': [],
                'tables': [],
                'image_py': [],
                'image_codeblock': [],
                'image_markdown': [],
                'image_mm': [],
                'image_text': [],
                'table_text': [table_text_filename],
                'table_py': [py_filename],
                'table_codeblock': [],
                'table_markdown': [md_filename],    
                'type': 'table'    
            })

            chunk_index += 1


    ingestion_pipeline_dict['chunks'] = chunks
    ingestion_pipeline_dict['table_text_files'] = table_text_files


    return ingestion_pipeline_dict

def extract_table_number(filename, verbose=False):
    match = re.search(r"chunk_\d+_table_(\d+).png", filename)
    if match:
        table_number = match.group(1)
        if verbose: print(f"Extracted table number: {table_number}")
    else:
        table_number = '0'

    return table_number

def create_table_doc_chunks_with_table_images(ingestion_pipeline_dict):

    chunks = ingestion_pipeline_dict.get('chunks', [])
    chunk_index = len(chunks)
    tables_folder = ingestion_pipeline_dict['tables_directory']
    tables = ingestion_pipeline_dict['table_images'] 

    for index, tc in enumerate(tables):
        chunks.append({
            'chunk_number':chunk_index+1, 
            'text_file': '',
            'full_chunk_text':'',
            'images': [],
            'tables': [tc],
            'chunk_image_path': tc,
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'chunk_table_number': extract_table_number(tc),
            'image_mm': [],
            'image_text': [],
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],    
            'type': 'table'    
        })

        chunk_index += 1

    ingestion_pipeline_dict['chunks'] = chunks

    return ingestion_pipeline_dict

def create_doc_chunks_with_doc_int_markdown(ingestion_pipeline_dict):

    chunks = []

    text_directory = ingestion_pipeline_dict['text_directory']
    images_directory = ingestion_pipeline_dict['images_directory'] 
    tables_directory = ingestion_pipeline_dict['tables_directory']

    images = ingestion_pipeline_dict['image_files']
    tables_dfs = ingestion_pipeline_dict['tables_py'] 
    tables_md = ingestion_pipeline_dict['tables_md'] 
    tables = ingestion_pipeline_dict['table_files'] 

    model_info = ingestion_pipeline_dict['models'][0]
    n_tokens = ingestion_pipeline_dict['chunk_size']
    overlap = ingestion_pipeline_dict['chunk_overlap']

    text_files = []
    image_text_files = []
    table_text_files = []


    try:
        # text_chunks = semantic_chunk_text_file(ingestion_pipeline_dict['full_text_file'])
        text_chunks = sentence_chunk_text_file(ingestion_pipeline_dict['full_text_file'],
                                               chunk_size=n_tokens,
                                               chunk_overlap=overlap
                                            )
    except:
        text_chunks = []
        logc("Warning Only", f"Could not perform Semantic Chunking of file {ingestion_pipeline_dict['full_text_file']}")

    chunk_index = 0

    for index, tc in enumerate(text_chunks):
        text_filename = os.path.join(text_directory, f"chunk_{index}.txt")
        write_to_file(tc, text_filename, 'w')
        chunks.append({
            'chunk_number':chunk_index+1, 
            'text_file': text_filename,
            'full_chunk_text':'',
            'images': [],
            'tables': [],
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'image_mm': [],
            'image_text': [],
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],
            'type': 'text'
        })

        text_files.append(text_filename)
        chunk_index += 1


    for index, tc in enumerate(tables):
        md_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.md")
        py_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.py")
        table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.txt")

        md_table = read_asset_file(md_filename)[0]
        table_text = read_asset_file(table_text_filename)[0]

        if get_token_count(md_table) > 2 * n_tokens:
            ## First Backup the original table
            md_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.md.backup")
            table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_0.txt.backup")
            write_to_file(md_table, md_filename, 'w')
            write_to_file(table_text, table_text_filename, 'w')
            
            ## Table too big needs to be broken into chunks
            md_chunks, header, summary = chunk_markdown_table(md_table, model_info, n_tokens = n_tokens, overlap = overlap)

            for md_index, md_chunk in enumerate(md_chunks):
                md_filename = os.path.join(tables_directory, f"chunk_{index}_table_{md_index}.md")
                table_text_filename = os.path.join(tables_directory, f"chunk_{index}_table_{md_index}.txt")
                table_text_files.append(table_text_filename)

                text_contents = f"Table Summary:\n{summary}\n\nTable:\n{md_chunk}"
                write_to_file(md_chunk, md_filename, 'w')
                write_to_file(text_contents, table_text_filename, 'w')

                chunks.append({
                    'chunk_number':chunk_index+1, 
                    'text_file': table_text_filename,
                    'full_chunk_text':'',
                    'images': [],
                    'tables': [],
                    'image_py': [],
                    'image_codeblock': [],
                    'image_markdown': [],
                    'image_mm': [],
                    'image_text': [],
                    'table_text': [table_text_filename],
                    'table_py': [py_filename],
                    'table_codeblock': [],
                    'table_markdown': [md_filename],    
                    'type': 'table'    
                })

                chunk_index += 1
        else:        
            table_text_files.append(table_text_filename)

            chunks.append({
                'chunk_number':chunk_index+1, 
                'text_file': table_text_filename,
                'full_chunk_text':'',
                'images': [],
                'tables': [],
                'image_py': [],
                'image_codeblock': [],
                'image_markdown': [],
                'image_mm': [],
                'image_text': [],
                'table_text': [table_text_filename],
                'table_py': [py_filename],
                'table_codeblock': [],
                'table_markdown': [md_filename],    
                'type': 'table'    
            })

            chunk_index += 1


    for index, tc in enumerate(images):   
        chunks.append({
            'chunk_number':chunk_index+1, 
            'text_file': '',
            'full_chunk_text':'',
            'chunk_image_path': tc,
            'images': [tc],
            'tables': [],
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'image_mm': [],
            'image_text': [],            
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],        
            'post_process_image_with_context': False, 
            'type': 'image'
        })

        chunk_index += 1

    ingestion_pipeline_dict['chunks'] = chunks

    ingestion_pipeline_dict['text_files'] = text_files
    ingestion_pipeline_dict['image_text_files'] = image_text_files
    ingestion_pipeline_dict['table_text_files'] = table_text_files


    return ingestion_pipeline_dict

def extract_doc_using_doc_int(ingestion_pipeline_dict):

    doc_path = ingestion_pipeline_dict['document_path'] 
    images_folder = ingestion_pipeline_dict['images_directory'] 
    tables_folder = ingestion_pipeline_dict['tables_directory']
    table_extraction_mode = ingestion_pipeline_dict.get('extract_docint_tables_mode', "Markdown")

    document_intelligence_client = DocumentIntelligenceClient(endpoint=DI_ENDPOINT, credential=AzureKeyCredential(DI_KEY))

    with open(doc_path, "rb") as f:
        poller = document_intelligence_client.begin_analyze_document(
            "prebuilt-layout", analyze_request=f, output_content_format=ContentFormat.MARKDOWN, content_type="application/octet-stream"
        )
    result: AnalyzeResult = poller.result()
    
    content = result['content']
    

    tables_txt = []
    tables_dfs = []
    tables_md = []
    table_count = 0
    table_bounds = []

    tables = []

    if table_extraction_mode == "Markdown":
        
        tables = extract_markdown_table(content)
        for table in tables:
            try:
                table = table[0]
                table_path_md = os.path.join(tables_folder, f'chunk_{table_count}_table_0.md')
                table_path = os.path.join(tables_folder, f'chunk_{table_count}_table_0.txt')
                table_path_py = os.path.join(tables_folder, f'chunk_{table_count}_table_0.py')

                write_to_file(table, table_path_md, 'w')
                write_to_file(table, table_path, 'w')
                df = extract_markdown_table_as_df(table)
                py_script = f"df_{generate_uuid_from_string(str(df.to_dict())).replace('-', '_')} = pd.DataFrame.from_dict({df.to_dict()})"
                write_to_file(py_script, table_path_py, 'w')

                tables_dfs.append(table_path_py)
                tables_md.append(table_path_md)
                tables_txt.append(table_path)
                table_count += 1
                
                logc("Extracting Markdown Table", table_path)

            except Exception as e:
                logc("Error extracting tables", f"Error extracting tables from text '{table[:50]}'. Error: {e}")

    elif table_extraction_mode == "JustExtract":
        if 'tables' in result: table_bounds = result['tables']
        pages_trkr = {}
        tbl_number = 0

        try:
            for bounding in table_bounds:
                page_number = bounding['boundingRegions'][0]['pageNumber']
                polygon = bounding['boundingRegions'][0]['polygon']
                polygons = [x['polygon'] for x in bounding['boundingRegions']]
                if page_number not in pages_trkr: 
                    pages_trkr[page_number] = 1
                    tbl_number = 0
                png_file = os.path.join(ingestion_pipeline_dict['chunks_as_images_directory'], f'chunk_{page_number}.png')
                target_filename = os.path.join(tables_folder, f'chunk_{page_number}_table_{tbl_number}.png')
                tbl_number+= 1
                extract_figure(png_file, polygons, target_filename)
                tables.append(target_filename)

                logc("Extracting Image Table with Bounds", target_filename)
        except Exception as e:
            print(f"extract_doc_using_doc_int::Error extracting table: {e}")
    
    else:
        logc("WARNING", f"Table extraction mode {table_extraction_mode} not supported. Please use 'Markdown' or 'JustExtract'.")


    images = []
    img_nums = {}

    ### WARNING: DOC INT DOES EXTRACT FIGURES FOR MS WORD DOCX 
    if (ingestion_pipeline_dict['extension'] == '.docx') or (ingestion_pipeline_dict['extension'] == '.doc'):
        logc("WARNING", f"Document Intelligence does not extract images from MS DOC or DOCX files. If images are important, please use the 'py-docx' mode to extract images from DOCX files.")

    image_bounds = []
    if 'figures' in result: image_bounds = result['figures']

    try:
        for bounding in image_bounds:
            page_number = bounding['boundingRegions'][0]['pageNumber']
            img_number = img_nums.get(page_number, 0)
            img_nums[page_number] = img_number

            polygons = [x['polygon'] for x in bounding['boundingRegions']]
            png_file = os.path.join(ingestion_pipeline_dict['chunks_as_images_directory'], f'chunk_{page_number}.png')
            target_filename = os.path.join(images_folder, f'chunk_{page_number}_image_{img_number}.png')
            img_nums[page_number] += 1
            extract_figure(png_file, polygons, target_filename)
            images.append(target_filename)

            logc("Extracting Image with Bounds", target_filename)

    except Exception as e:
        logc(f"extract_doc_using_doc_int::Error extracting image: {e}")


    ingestion_pipeline_dict['full_text'] = result['content']
    write_to_file(result['content'], ingestion_pipeline_dict['full_text_file'], 'w')

    ingestion_pipeline_dict['table_images'] = tables
    ingestion_pipeline_dict['table_files'] = tables_txt
    ingestion_pipeline_dict['tables_py'] = tables_dfs
    ingestion_pipeline_dict['tables_md'] = tables_md
    ingestion_pipeline_dict['image_files'] = images
    
    return ingestion_pipeline_dict

def extract_docx_using_py_docx(ingestion_pipeline_dict):
    """
    This function modifies the 'ingestion_pipeline_dict' dictionary in the following ways:

    Reads from 'ingestion_pipeline_dict':
    - 'document_path': The path to the .docx document to be processed.
    - 'images_directory': The directory path where extracted images will be stored.
    - 'tables_directory': The directory path where extracted tables will be stored as Markdown, plain text, and Python files.

    Writes to 'ingestion_pipeline_dict':
    - 'full_text': A concatenated string of all the text extracted from the document.
    - 'full_text_file': The path to a file where the full concatenated text is saved. (Note: The actual file path is not provided in the input dictionary but is assumed to be dynamically determined within the function.)
    - 'tables_py': A list of file paths to the generated Python scripts that recreate the tables using pandas.
    - 'tables_md': A list of file paths to the Markdown files containing the tables.
    - 'image_files': A list of file paths to the extracted images.
    - 'table_files': A list of file paths to the plain text files containing the tables.
    """
    doc_path = ingestion_pipeline_dict['document_path'] 
    images_folder = ingestion_pipeline_dict['images_directory'] 
    tables_folder = ingestion_pipeline_dict['tables_directory']

    doc = Document(doc_path)
    all_text = []
    tables = []
    tables_dfs = []
    tables_md = []
    images = []
    image_count = 0
    table_count = 0
    
    # Ensure the images_folder exists
    if not os.path.exists(images_folder):
        os.makedirs(images_folder)
    
    # Extract all text
    for para in doc.paragraphs:
        all_text.append(para.text)
    
    # Extract all tables
    for table in doc.tables:
        try:
            headers = [cell.text for cell in table.rows[0].cells]
            data = []
            for row in table.rows[1:]:  # Skip header row
                data.append([cell.text for cell in row.cells])
            df = pd.DataFrame(data, columns=headers)
            table_path, table_path_md, table_path_py = write_df_to_files(df, tables_folder, table_count)
            tables_dfs.append(table_path_py)
            tables_md.append(table_path_md)
            tables.append(table_path)
            table_count += 1
        except Exception as e:
            logc(f"Couldnt extract table from file: {doc_path}")
    
    # Extract all images
    for rel in doc.part.rels.values():
        try:
            if "image" in rel.target_ref:
                img = rel.target_part.blob
                image_path = os.path.join(images_folder, f'chunk_{image_count}_image_0.jpg')
                with open(image_path, 'wb') as img_file:
                    img_file.write(img)
                image_count += 1
                images.append(image_path)
        except Exception as e:
            logc(f"Couldnt extract image from file: {doc_path}")
    
    concatenated_text = '\n'.join(all_text)

    ingestion_pipeline_dict['full_text'] = concatenated_text
    write_to_file(concatenated_text, ingestion_pipeline_dict['full_text_file'], 'w')

    ingestion_pipeline_dict['tables_py'] = tables_dfs
    ingestion_pipeline_dict['tables_md'] = tables_md
    ingestion_pipeline_dict['image_files'] = images
    ingestion_pipeline_dict['table_files'] = tables

    return ingestion_pipeline_dict
    
def extract_audio_using_whisper(ingestion_pipeline_dict):

    ### generate concatenated_text
    concatenated_text = ""
    ingestion_pipeline_dict['full_text'] = concatenated_text
    write_to_file(concatenated_text, ingestion_pipeline_dict['full_text_file'], 'w')

    return ingestion_pipeline_dict

def create_pdf_chunks(ingestion_pipeline_dict):
    document_file_path = ingestion_pipeline_dict['document_path']
    password = ingestion_pipeline_dict.get('password', None)

    # Open the PDF file
    pdf_document = fitz.open(document_file_path)
    full_basename = os.path.basename(document_file_path)

    if password is not None: 
        r = pdf_document.authenticate(password)
        if r == 0: raise ValueError("Password is incorrect.")
        filename = document_file_path + '.decrypted.pdf'
        pdf_document.save(filename)
        logc(f"Ingestion Stage of {full_basename}- Info", f"Opened the file with the password. Status: {r}")

    logc(f"Ingestion Stage of {full_basename} - Info", f"PDF File with num chunks -> {len(pdf_document)}")

    ingestion_pipeline_dict['num_chunks'] = len(pdf_document)
    ingestion_pipeline_dict['document_file_path'] = document_file_path
    ingestion_pipeline_dict['pdf_document'] = pdf_document
    
    if 'chunks' in ingestion_pipeline_dict:
        for index, chunk in enumerate(pdf_document):
            ingestion_pipeline_dict['chunks'][index]['chunk'] = chunk
    else:
        ingestion_pipeline_dict['chunks'] = [{
            'chunk':chunk, 
            'text_file': '',
            'chunk_number':index+1, 
            'full_chunk_text':'',
            'images': [],
            'tables': [],
            'image_py': [],
            'image_codeblock': [],
            'image_markdown': [],
            'image_mm': [],
            'image_text': [],
            'table_text': [],
            'table_py': [],
            'table_codeblock': [],
            'table_markdown': [],   
            'type': 'page',
            'post_process_image_with_context': True,     
        } for index, chunk in enumerate(pdf_document)]

    return ingestion_pipeline_dict

def harvest_code_from_text(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):

    text_filename = chunk_dict['text_file']
    py_file = replace_extension(text_filename, '.py')
    codeblock_file = replace_extension(text_filename, '.codeblock')
    markdown_file = replace_extension(text_filename, '.md')

    data = read_asset_file(text_filename)[0]

    if data == '': 
        logc(f"Harvesting Code: Empty text file: {text_filename}")
        return []

    code_harvesting_prompt = code_harvesting_from_text.format(text=data, random_block_id=str(uuid.uuid4())[:8])

    messages = []
    messages.append({"role": "system", "content": "You are a helpful AI assistant who specializes in Python code generation. You help users answer their queries based on the information supplied below." })     
    messages.append({"role": "user", "content": code_harvesting_prompt})     

    if not os.path.exists(py_file):
        try:
            client = AzureOpenAI(
                azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" , 
                api_key= model_info['AZURE_OPENAI_KEY'],  
                api_version= AZURE_OPENAI_API_VERSION,
            )

            result = get_chat_completion(messages, model=model_info['AZURE_OPENAI_MODEL'], client = client)
            response = result.choices[0].message.content
            # print(f"Harvested Code from chunk {extract_chunk_number(text_filename)}:", response)

            py_code = extract_code(response)
            codeblock = "```python\n" + py_code + "\n```"
            markdown_table = extract_markdown(response)

            write_to_file(codeblock, codeblock_file)
            write_to_file(py_code, py_file)
            write_to_file(markdown_table, markdown_file)

            chunk_dict['codeblock_file'] = codeblock_file
            chunk_dict['py_file'] = py_file
            chunk_dict['markdown_file'] = markdown_file

            return [{'codeblock_file':codeblock_file, 'py_file':py_file, 'markdown_file':markdown_file}]
        except Exception as e:
            print("harvest_code_from_text Error:", e)
            return []
    else:
        if os.path.exists(codeblock_file): chunk_dict['codeblock_file'] = codeblock_file
        if os.path.exists(py_file): chunk_dict['py_file'] = py_file
        if os.path.exists(markdown_file): chunk_dict['markdown_file'] = markdown_file

        return [{'codeblock_file':codeblock_file, 'py_file':py_file, 'markdown_file':markdown_file}]

def harvest_code(ingestion_pipeline_dict):
    harvested_code, _ = execute_multithreaded_funcs(harvest_code_from_text, ingestion_pipeline_dict)
    ingestion_pipeline_dict['harvested_code'] = harvested_code
    for code_dict in harvested_code:
        code = read_asset_file(code_dict['py_file'])[0]
        write_to_file(code + '\n\n', ingestion_pipeline_dict['master_py_file'], mode='a')
        ingestion_pipeline_dict['py_files'].append(code_dict['py_file'])
        ingestion_pipeline_dict['codeblock_files'].append(code_dict['codeblock_file'])
        ingestion_pipeline_dict['markdown_files'].append(code_dict['markdown_file'])

    return ingestion_pipeline_dict

def pdf_extract_high_res_chunk_images(ingestion_pipeline_dict):

    high_res_chunk_images = []
    chunks_as_images_directory = ingestion_pipeline_dict['chunks_as_images_directory']

    for index, chunk_dict in enumerate(ingestion_pipeline_dict['chunks']):
        # print(chunk_dict)
        try:
            chunk = chunk_dict['chunk']
            chunk_number = chunk_dict['chunk_number']

            chunk_pix = chunk.get_pixmap(dpi=300)
            cropbox = chunk.cropbox
            chunk.set_cropbox(chunk.mediabox)
            image_filename = f'chunk_{chunk_number}.png'
            image_path = os.path.join(chunks_as_images_directory, image_filename)
            chunk_pix.save(image_path)
            high_res_chunk_images.append(image_path)
            
            image_filename_lowres = f'chunk_{chunk_number}_lowres.png'
            image_path_lowres = os.path.join(chunks_as_images_directory, image_filename_lowres)
            chunk_pix = chunk.get_pixmap(dpi=150)
            chunk_pix.save(image_path_lowres)

            chunk_dict['chunk_image_path'] = image_path
            # chunk_dict['cropbox'] = cropbox
            # chunk_dict['a4_or_slide'] = 'a4' if cropbox[2] < cropbox[3] else 'slide'
        except Exception as e: 
            print(f"Exception in the {index}-th chunk.\nException:{e}\ntraceback_print: {str(traceback.print_exc())}\ntraceback_format: {str(traceback.format_exc())}") 
            

    ingestion_pipeline_dict['high_res_chunk_images']  = high_res_chunk_images

    return ingestion_pipeline_dict

def process_images_with_GPT4V(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    
    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    image_path = chunk_dict['chunk_image_path']
    images_directory = ingestion_pipeline_dict['images_directory']
    print(f"Processing image {index} on chunk {chunk_number} with model {model_info['AZURE_OPENAI_RESOURCE']}")
    image_filename = None
    detected_filename = replace_extension(image_path, '.detected.txt')

    if not os.path.exists(detected_filename):
        try:
            count, description, _ = get_asset_explanation_gpt4v(image_path, None, gpt4v_prompt = detect_num_of_diagrams_prompt, with_context=False, extension='dont_save', model_info=model_info)
            write_to_file(count, detected_filename, 'w')
            image_count = int(count)
            print(f"Number of Images Detected in chunk number {chunk_number} : {count}.")
        except Exception as e:
            print(f"Error in image detection: {e}")
    else:
        try:
            image_count = int(read_asset_file(detected_filename)[0])
        except:
            image_count = 0 
            print(f"Error reading image count from file: {detected_filename}")


    if image_count > 0:
        print("Image Detection", f"{bc.OKBLUE}Image Detection Status on chunk {chunk_number}: {bc.OKGREEN}OK - Detected {image_count} images.{bc.ENDC}")
        image_filename = os.path.join(images_directory, f'chunk_{chunk_number}_image_{index+1}.jpg')
        shutil.copyfile(image_path, image_filename)
        print(f"Saved Image {image_count+1} on chunk {chunk_number} to '{image_filename}'")
        chunk_dict['images'] = [image_filename]
        return [image_filename]
    
    return []

def process_images_with_PDF(ingestion_pipeline_dict, chunk_dict):
    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    images_directory = ingestion_pipeline_dict['images_directory']
    chunk = chunk_dict['chunk']
    pdf_document = ingestion_pipeline_dict['pdf_document']
    image_files = []

    for img_index, img in enumerate(chunk.get_images()):
        xref = img[0]
        base_image = pdf_document.extract_image(xref)
        pix = fitz.Pixmap(pdf_document, xref)
        pix = fitz.Pixmap(fitz.csRGB, pix)                
        image_filename = os.path.join(images_directory, f'chunk_{chunk_number}_image_{img_index+1}.png')
        pix.save(image_filename, 'PNG')
        image_files.append(image_filename)

    return image_files

def execute_multithreaded_funcs(func, ingestion_pipeline_dict, args = None):
    return_array = []

    num_threads = ingestion_pipeline_dict['num_threads']
    # num_chunks = ingestion_pipeline_dict['num_chunks']
    num_chunks = len(ingestion_pipeline_dict['chunks'])
    rounds = math.ceil(num_chunks / num_threads)
    last_round = num_chunks % num_threads
    chunks = ingestion_pipeline_dict['chunks']

    if (args is not None) and (args.get('vision_models', False)):
        models = ingestion_pipeline_dict['vision_models']
    else:
        models = ingestion_pipeline_dict['models']

    logc(f"Last Round Remainder: {last_round} chunks. Num chunks: {num_chunks}. Num Threads: {num_threads}. Rounds: {rounds}.")

    for r in range(rounds):
        list_pipeline_dict = [ingestion_pipeline_dict] * num_threads
        list_chunk_dict = chunks[r*num_threads:(r+1)*num_threads]
        list_index = [x for x in range(r*num_threads+1,(r+1)*num_threads+1)]
        list_args = [args] * num_threads

        if (last_round > 0) and (r == rounds - 1): # last round
            list_pipeline_dict = list_pipeline_dict[:last_round]
            list_chunk_dict = list_chunk_dict[:last_round]
            list_index = list_index[:last_round]

        logc("Processing...", f"Round {r+1} of {rounds} with {len(list_chunk_dict)} chunks and {num_threads} threads.")
        results = pool.starmap(func,  zip(list_pipeline_dict, list_chunk_dict, models, list_index, list_args))
        for i in results: return_array.extend(i)

    return return_array, ingestion_pipeline_dict

def pdf_extract_images(ingestion_pipeline_dict):
    image_files = []

    extract_images_mode = ingestion_pipeline_dict['extract_images_mode']
    args = {'vision_models': True}

    if extract_images_mode == "GPT":
        image_files, _ = execute_multithreaded_funcs(process_images_with_GPT4V, ingestion_pipeline_dict, args=args)

    elif extract_images_mode == "PDF":
        for chunk_dict in ingestion_pipeline_dict['chunks']:
            chunk_image_files = process_images_with_PDF(ingestion_pipeline_dict, chunk_dict)
            chunk_dict['images'] = chunk_image_files
            image_files += chunk_image_files

    else:
        raise ValueError(f"Unsupported extract_images_mode: {extract_images_mode}")

    ingestion_pipeline_dict['image_files'] = image_files
    return ingestion_pipeline_dict



def generate_tags_with_GPT4(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    
    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    text_file = chunk_dict['text_file']
    text_directory = ingestion_pipeline_dict['text_directory']
    azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" 
    print(f"GPT4 Tags - Extraction - Processing text {index} on chunk {chunk_number} using {model_info['AZURE_OPENAI_MODEL']} and endpoint {azure_endpoint}")
    # tags_filename = os.path.join(text_directory, f'chunk_{chunk_number}.tags.txt')
    tags_filename = replace_extension(text_file, '.tags.txt')
    
    try:
        client = AzureOpenAI(
            azure_endpoint =  azure_endpoint, 
            api_key= model_info['AZURE_OPENAI_KEY'],  
            api_version= AZURE_OPENAI_API_VERSION,
        )

        if not os.path.exists(tags_filename):
            text = read_asset_file(text_file)[0]
            print(f"GPT4 Tags - Post-Processing: Generating tags for chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")
            optimized_tag_list = generate_tag_list(text, model = model_info['AZURE_OPENAI_MODEL'], client = client)
            write_to_file(optimized_tag_list, tags_filename ,"w")
            chunk_dict['tags_file'] = tags_filename
            print(f"GPT4 Tags - Post-Processing: Tags processed in chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")

        return [tags_filename]

    except Exception as e:
        print(f"generate_tags_with_GPT4::Error in text processing in model {model_info['AZURE_OPENAI_RESOURCE']}:\nFor text file: {text_file}\n{e}")

    return []

def generate_tags_for_text(ingestion_pipeline_dict):
    tags_files = []

    ingestion_pipeline_dict_ret = copy.deepcopy(ingestion_pipeline_dict)
    ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if (rd['text_file'] != '') and (rd['type'] == 'text')]

    tags_filenames, _ = execute_multithreaded_funcs(generate_tags_with_GPT4, ingestion_pipeline_dict_ret)
    tags_files.extend(tags_filenames)

    for index1, rd in enumerate(ingestion_pipeline_dict_ret['chunks']):
        for index2, r in enumerate(ingestion_pipeline_dict['chunks']):
            if rd['chunk_number'] == r['chunk_number']:
                ingestion_pipeline_dict['chunks'][index2] = copy.deepcopy(rd)

    ingestion_pipeline_dict['tags_files'] = tags_files

    return ingestion_pipeline_dict

def generate_tags_for_all_chunks(ingestion_pipeline_dict):
    tags_files = []

    ingestion_pipeline_dict_ret = copy.deepcopy(ingestion_pipeline_dict)
    ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if (rd['text_file'] != '') and (read_asset_file(rd['text_file'])[0] != '')]

    # print("ingestion_pipeline_dict_ret", ingestion_pipeline_dict_ret)
    tags_filenames, _ = execute_multithreaded_funcs(generate_tags_with_GPT4, ingestion_pipeline_dict_ret)
    tags_files.extend(tags_filenames)

    for index1, rd in enumerate(ingestion_pipeline_dict_ret['chunks']):
        for index2, r in enumerate(ingestion_pipeline_dict['chunks']):
            if rd['chunk_number'] == r['chunk_number']:
                ingestion_pipeline_dict['chunks'][index2] = copy.deepcopy(rd)

    ingestion_pipeline_dict['tags_files'] = tags_files

    return ingestion_pipeline_dict



def generate_analsysis_with_GPT4(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    
    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    text_file = chunk_dict['text_file']
    text_directory = ingestion_pipeline_dict['text_directory']
    text_summary_path = replace_extension(ingestion_pipeline_dict['full_text_file'], '.summary.txt')
    original_document_filename = ingestion_pipeline_dict['original_document_filename']
    azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" 
    print(f"GPT4 Tags - Extraction - Processing text {index} on chunk {chunk_number} using {model_info['AZURE_OPENAI_MODEL']} and endpoint {azure_endpoint}")
    analysis_filename = replace_extension(text_file, '.analysis.txt')
    
    try:
        client = AzureOpenAI(
            azure_endpoint =  azure_endpoint, 
            api_key= model_info['AZURE_OPENAI_KEY'],  
            api_version= AZURE_OPENAI_API_VERSION,
        )

        if not os.path.exists(analysis_filename):
            text_summary = read_asset_file(text_summary_path)[0]
            
            text = read_asset_file(text_file)[0]
            print(f"GPT4 Analysis - Post-Processing: Generating analysis for chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")
            # text_tokens = get_token_count(text)
            # prompt_template_tokens = get_token_count(chunk_analysis_template)
            # full_text = limit_token_count(text_summary, limit = (FULL_TEXT_TOKEN_LIMIT - text_tokens - prompt_template_tokens))
            prompt = chunk_analysis_template.format(text_summary=text_summary, text_chunk=text, chunk_number=chunk_number, filename=original_document_filename)
            analysis = ask_LLM(prompt)
            write_to_file(analysis, analysis_filename ,"w")
            chunk_dict['analysis_file'] = analysis_filename
            print(f"GPT4 Analysis - Post-Processing: Analysis processed in chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")

        return [analysis_filename]

    except Exception as e:
        print(f"generate_analsysis_with_GPT4::Error in text processing in model {model_info['AZURE_OPENAI_RESOURCE']}:\nFor text file: {text_file}\n{e}")

    return []

def generate_analysis_for_text(ingestion_pipeline_dict):
    analysis_files = []

    ingestion_pipeline_dict_ret = copy.deepcopy(ingestion_pipeline_dict)
    # ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if (rd['text_file'] != '') and (rd['type'] == 'text')]
    ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if (rd['text_file'] != '')]

    analysis_filenames, _ = execute_multithreaded_funcs(generate_analsysis_with_GPT4, ingestion_pipeline_dict_ret)
    analysis_files.extend(analysis_filenames)

    for index1, rd in enumerate(ingestion_pipeline_dict_ret['chunks']):
        for index2, r in enumerate(ingestion_pipeline_dict['chunks']):
            if rd['chunk_number'] == r['chunk_number']:
                ingestion_pipeline_dict['chunks'][index2] = copy.deepcopy(rd)

    ingestion_pipeline_dict['analysis_files'] = analysis_files

    return ingestion_pipeline_dict

def process_text_with_GPT4(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    
    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    text_file = chunk_dict['text_file']
    text_directory = ingestion_pipeline_dict['text_directory']
    azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" 
    print(f"GPT4 Text - Extraction - Processing text {index} on chunk {chunk_number} using {model_info['AZURE_OPENAI_MODEL']} and endpoint {azure_endpoint}")
    original_text_filename = os.path.join(text_directory, f'chunk_{chunk_number}.original.txt')
    processed_text_filename = os.path.join(text_directory, f'chunk_{chunk_number}.processed.txt')

    try:
        client = AzureOpenAI(
            azure_endpoint =  azure_endpoint, 
            api_key= model_info['AZURE_OPENAI_KEY'],  
            api_version= AZURE_OPENAI_API_VERSION,
        )

        if not os.path.exists(processed_text_filename):
            messages = []
            messages.append({"role": "system", "content": "You are a helpful assistant that helps the user by generating high quality code to answer the user's questions."})     
            messages.append({"role": "user", "content": process_extracted_text_prompt.format(text=read_asset_file(text_file)[0], markdown="No Markdown available.")})     

            result = get_chat_completion(messages, model=model_info['AZURE_OPENAI_MODEL'], client = client)     
            response = result.choices[0].message.content

            shutil.copyfile(text_file, original_text_filename)
            write_to_file(response, text_file, 'w')
            write_to_file(response, processed_text_filename)
            chunk_dict['original_text'] = original_text_filename
            chunk_dict['processed_text'] = processed_text_filename

            # time.sleep(2)
            print(f"GPT4 Text - Post-Processing: Generating tags for chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")
            optimized_tag_list = generate_tag_list(response, model = model_info['AZURE_OPENAI_MODEL'], client = client)
            write_to_file(optimized_tag_list, replace_extension(text_file, '.tags.txt'))

            print(f"GPT4 Text - Post-Processing: Text processed in chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")

        chunk_dict['original_text'] = original_text_filename
        chunk_dict['processed_text'] = processed_text_filename
        shutil.copyfile(processed_text_filename, text_file)
        return [original_text_filename]

    except Exception as e:
        print(f"process_text_with_GPT4::Error in text processing in model {model_info['AZURE_OPENAI_RESOURCE']}:\nFor text file: {text_file}\n{e}")

    return []

def pdf_extract_text(ingestion_pipeline_dict):
    text_files = []
    original_text_files = []
    text_directory = ingestion_pipeline_dict['text_directory']

    extract_text_mode = ingestion_pipeline_dict['extract_text_mode']
    models = ingestion_pipeline_dict['models']
    num_threads = ingestion_pipeline_dict['num_threads']

    for chunk_dict in ingestion_pipeline_dict['chunks']:
        #### 4 SAVE PDF chunks AS TEXT
        chunk = chunk_dict['chunk']
        chunk_number = chunk_dict['chunk_number']
        text = chunk.get_text()
        # Define the filename for the current chunk

        text_filename = os.path.join(text_directory, f"chunk_{chunk_number}.txt")
        # Save the text to a file
        with open(text_filename, 'w', encoding='utf-8') as file:
            file.write(text)

        text_files.append(text_filename)
        chunk_dict['text_file'] = text_filename

    if extract_text_mode == "GPT":
        original_text_files, _ = execute_multithreaded_funcs(process_text_with_GPT4, ingestion_pipeline_dict)

    ingestion_pipeline_dict['text_files'] = text_files
    ingestion_pipeline_dict['original_text_files'] = original_text_files

    for chunk in ingestion_pipeline_dict['chunks']:
        text = read_asset_file(chunk['text_file'])[0]
        write_to_file(text + '\n\n', ingestion_pipeline_dict['full_text_file'], mode='a')

    return ingestion_pipeline_dict

def extract_table_from_image(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    #### 2 DETECT AND SAVE TABLES
    table_number = chunk_dict.get('chunk_table_number', 0)
    chunk_number = chunk_dict['chunk_number']
    image_path = chunk_dict['chunk_image_path']
    tables_directory = ingestion_pipeline_dict['tables_directory']
    table_filename = os.path.join(tables_directory, f"chunk_{chunk_number}_table_{table_number}.png")
    detected_filename = replace_extension(table_filename, '.detected.txt')

    if not os.path.exists(detected_filename):
        try:
            count, description, _ = get_asset_explanation_gpt4v(image_path, None, gpt4v_prompt = detect_num_of_tables_prompt, with_context=False, extension='dont_save', model_info=model_info)
            print(f"Table Detection {count} in chunk {chunk_number}")
            table_count = int(count)
            status = f"OK - Detected {table_count} tables."
            write_to_file(count, detected_filename, 'w')

        except Exception as e:
            print(f"Error in table detection: {e}")
            status = f"Error Detecting number of tables. Exception: {e}"
            table_count = 0

        print(f"{bc.OKBLUE}Table Detection Status on chunk {chunk_number}: {bc.OKGREEN}{status}{bc.ENDC}")

    else:
        try:
            table_count = int(read_asset_file(detected_filename)[0])
        except:
            table_count = 0 
            print(f"Error reading table count from file: {detected_filename}")  

    if table_count > 0:
        shutil.copyfile(image_path, table_filename)
        print(f"Saved table {table_number} on chunk {chunk_number} to '{table_filename}'")
        chunk_dict['tables'] = [table_filename]
        return [table_filename]
    return []

def extract_tables_from_images(ingestion_pipeline_dict):
    args = {'vision_models': True}

    tables, _ = execute_multithreaded_funcs(extract_table_from_image, ingestion_pipeline_dict, args=args)
    ingestion_pipeline_dict['tables'] = tables
    return ingestion_pipeline_dict

def post_process_images(ingestion_pipeline_dict):

    extract_text_from_images = ingestion_pipeline_dict['extract_text_from_images']
    args = {'extract_text_from_images':extract_text_from_images, 'vision_models': True}

    ingestion_pipeline_dict_ret = copy.deepcopy(ingestion_pipeline_dict)
    ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if len(rd['images']) > 0]

    image_proc_files, ingestion_pipeline_dict_ret = execute_multithreaded_funcs(post_process_chunk_images, ingestion_pipeline_dict_ret, args=args)

    for index1, rd in enumerate(ingestion_pipeline_dict_ret['chunks']):
        for index2, r in enumerate(ingestion_pipeline_dict['chunks']):
            if rd['chunk_number'] == r['chunk_number']:
                ingestion_pipeline_dict['chunks'][index2] = copy.deepcopy(rd)
    
    # print("\n\nChunk_dicts", json.dumps(ingestion_pipeline_dict_ret['chunks'], indent=4))
    # print("\n\nChunk_dicts", json.dumps(ingestion_pipeline_dict['chunks'], indent=4))
    ingestion_pipeline_dict['image_proc_files'] = image_proc_files

    for image_dict in image_proc_files:
        for f in image_dict['image_py']:
            code = read_asset_file(f)[0]
            write_to_file(code + '\n\n', ingestion_pipeline_dict['master_py_file'], mode='a')

        ingestion_pipeline_dict['py_files'].extend(image_dict['image_py'])
        ingestion_pipeline_dict['codeblock_files'].extend(image_dict['image_codeblock'])
        ingestion_pipeline_dict['markdown_files'].extend(image_dict['image_markdown'])
        ingestion_pipeline_dict['mermaid_files'].extend(image_dict['image_mm'])
        ingestion_pipeline_dict['image_text_files'].extend(image_dict['image_text'])

    return ingestion_pipeline_dict



def post_process_chunk_images(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    
    if args is not None:
        extract_text_from_images = args.get('extract_text_from_images', True)
    else:
        extract_text_from_images = True

    image_count = 0
    chunk_number = chunk_dict['chunk_number']
    image_path = chunk_dict['chunk_image_path']
    chunk_text_file = chunk_dict['text_file']
    master_text_file = ingestion_pipeline_dict['full_text_file']
    images_directory = ingestion_pipeline_dict['images_directory']
    document_path = ingestion_pipeline_dict['document_path']
    print(f"Post-Processing image {index} on chunk {chunk_number} using model {model_info['AZURE_OPENAI_RESOURCE']}")
    image_filename = None
    image_py_files = []
    image_codeblock_files = []
    image_mm_files = []
    image_text_files = []
    image_markdown = []

    client = AzureOpenAI(
        azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" , 
        api_key= model_info['AZURE_OPENAI_KEY'],  
        api_version= AZURE_OPENAI_API_VERSION,
    )

    if extract_text_from_images:
         image_description_prompt_modified = image_description_prompt  + extract_text_from_images_prompt
    else:
        image_description_prompt_modified = image_description_prompt

    print(f"Chunk Dict Images: {chunk_dict['images']}")

    with_context = chunk_dict.get('post_process_image_with_context', False)

    for image in chunk_dict['images']:
        
        if not os.path.exists(replace_extension(image, '.tags.txt')):
            text, text_filename, _ = get_asset_explanation_gpt4v(image, document_path, gpt4v_prompt = image_description_prompt_modified, with_context=with_context,  extension='.txt', model_info=model_info)

            mrkdwn = extract_markdown(text)
            if mrkdwn != "":
                code_filename = text_filename.replace('.txt', '.md')
                write_to_file(mrkdwn, code_filename)


            ocr_text = extract_extracted_text(text)
            if (ocr_text != "") and (get_token_count(ocr_text) > 10):
                messages = []
                messages.append({"role": "system", "content": "You are a helpful assistant that helps the user by generating high quality code to answer the user's questions."})     
                messages.append({"role": "user", "content": process_extracted_text_prompt.format(text=ocr_text, markdown=mrkdwn)})     
                result = get_chat_completion(messages, model=model_info['AZURE_OPENAI_MODEL'], client = client)     
                response = result.choices[0].message.content

                text = remove_extracted_text(text) + "\n\n**Extracted Text:**\n" + response
                write_to_file(text, text_filename, 'w')

            py_code = extract_code(text)
            if py_code != "":
                code_filename = text_filename.replace('.txt', '.py')
                write_to_file(py_code, code_filename)
                image_py_files.append(code_filename)
                codeblock = "```python\n" + py_code + "\n```"
                block_filename = text_filename.replace('.txt', '.codeblock')
                write_to_file(codeblock, block_filename)
                image_codeblock_files.append(block_filename)

            mm_code = extract_mermaid(text)
            if mm_code != "":
                code_filename = text_filename.replace('.txt', '.mermaid')
                write_to_file(mm_code, code_filename)
                image_mm_files.append(code_filename)



            image_text_files.append(text_filename)

            # write_to_file(f'\n\n\n#### START OF DESCRIPTION OF IMAGE {index}\n' + remove_code(text) + '\n#### END OF DESCRIPTION OF IMAGE\n\n', chunk_text_file, mode='a')
            write_to_file(remove_code(text), text_filename, 'w')
            write_to_file(f'\n\n\n#### START OF DESCRIPTION OF IMAGE {index}\n' + remove_code(text) + f'\n#### END OF DESCRIPTION OF IMAGE {index}\n\n', master_text_file, mode='a')
            # write_to_file(remove_code(text) + '\n\n', master_text_file, mode='a')

            # time.sleep(2)
            optimized_tag_list = generate_tag_list(remove_code(text), model = model_info['AZURE_OPENAI_MODEL'], client = client)
            write_to_file(optimized_tag_list, replace_extension(text_filename, '.tags.txt'))

        else:
            print(f"Image Tags File Already Exists for file {image}")
            text_filename = replace_extension(image, '.txt')
            code_filename = text_filename.replace('.txt', '.py')
            if os.path.exists(code_filename): image_py_files.append(code_filename)
            block_filename = text_filename.replace('.txt', '.codeblock')
            if os.path.exists(block_filename): image_codeblock_files.append(block_filename)
            mm_filename = text_filename.replace('.txt', '.mermaid')
            if os.path.exists(mm_filename): image_mm_files.append(mm_filename)
            mrkdwn_filename = text_filename.replace('.txt', '.md')
            if os.path.exists(mrkdwn_filename): image_markdown.append(mrkdwn_filename)
            image_text_files.append(text_filename)



    print(f"Post-Processing: Image processed in chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")
    chunk_dict['image_py'] = image_py_files
    chunk_dict['image_codeblock'] = image_codeblock_files
    chunk_dict['image_mm'] = image_mm_files
    chunk_dict['image_text'] = image_text_files
    chunk_dict['image_markdown'] = image_markdown
    chunk_dict['text_file'] = text_filename

    # print("Chunk_dict", json.dumps(chunk_dict, indent=4))
    return [{'image_py':image_py_files, 'image_codeblock':image_codeblock_files, 'image_mm':image_mm_files, 'image_text':image_text_files, 'image_markdown':image_markdown}]

def post_process_tables(ingestion_pipeline_dict):
    ingestion_pipeline_dict_ret = copy.deepcopy(ingestion_pipeline_dict)
    # logc("\n\nAssets - Before Deletion", ingestion_pipeline_dict_ret['chunks'])
    ingestion_pipeline_dict_ret['chunks'] = [rd for rd in ingestion_pipeline_dict_ret['chunks'] if len(rd['tables']) > 0]
    # logc("\n\nAssets - After Deletion", ingestion_pipeline_dict_ret['chunks'])

    args = {'vision_models': True}

    table_proc_files, ingestion_pipeline_dict_ret = execute_multithreaded_funcs(post_process_chunk_table, ingestion_pipeline_dict_ret, args=args)

    # logc("\n\nRet Assets - After Processing", ingestion_pipeline_dict_ret['chunks'])

    for index1, rd in enumerate(ingestion_pipeline_dict_ret['chunks']):
        for index2, r in enumerate(ingestion_pipeline_dict['chunks']):
            if rd['chunk_number'] == r['chunk_number']:
                ingestion_pipeline_dict['chunks'][index2] = copy.deepcopy(rd)

    # logc("\n\nFull Assets - After Processing", ingestion_pipeline_dict['chunks'])                

    for table_dict in table_proc_files:
        for f in table_dict['table_py']:
            code = read_asset_file(f)[0]
            write_to_file(code + '\n\n', ingestion_pipeline_dict['master_py_file'], mode='a')
        ingestion_pipeline_dict['py_files'].extend(table_dict['table_py'])
        ingestion_pipeline_dict['codeblock_files'].extend(table_dict['table_codeblock'])
        ingestion_pipeline_dict['markdown_files'].extend(table_dict['table_markdown'])
        ingestion_pipeline_dict['table_text_files'].extend(table_dict['table_text'])

    return ingestion_pipeline_dict

def post_process_chunk_table(ingestion_pipeline_dict, chunk_dict, model_info = None, index = 0, args = None, verbose = False):
    chunk_number = chunk_dict['chunk_number']
    image_path = chunk_dict['chunk_image_path']
    chunk_text_file = chunk_dict['text_file']
    master_text_file = ingestion_pipeline_dict['full_text_file']
    tables_directory = ingestion_pipeline_dict['tables_directory']
    document_path = ingestion_pipeline_dict['document_path']


    print(f"Post-Processing table {index} on chunk {chunk_number}")
    table_text_files = []
    table_code_text_filenames = []
    table_code_py_filenames = []
    table_markdown_filenames = []

    client = AzureOpenAI(
        azure_endpoint =  f"https://{model_info['AZURE_OPENAI_RESOURCE']}.openai.azure.com" , 
        api_key= model_info['AZURE_OPENAI_KEY'],  
        api_version= AZURE_OPENAI_API_VERSION,
    )

    for table in chunk_dict['tables']:
        if not os.path.exists(replace_extension(table, '.tags.txt')):
            text, text_filename, _ = get_asset_explanation_gpt4v(table, document_path, gpt4v_prompt = image_description_prompt, with_context=True,  extension='.txt', model_info=model_info)

            markdown = extract_markdown(text)
            if markdown == "":
                markdown, markdown_filename, _ = get_asset_explanation_gpt4v(table, document_path, gpt4v_prompt = table_markdown_description_prompt, with_context=True, extension='.md', model_info=model_info)
            else: 
                markdown_filename = text_filename.replace('.txt', '.md')
                write_to_file(markdown, markdown_filename, 'w')

            code_execution_success = False
            temperature = 0.2
            retries = 0
            prompt_extension = ""

            code = extract_code(text)
            if code != "":
                code_filename = text_filename.replace('.txt', '.py')
                code_text_filename = text_filename.replace('.txt', '.codeblock')
                codeblock = "```python\n" + code + "\n```"
                write_to_file(code, code_filename, 'w')
                write_to_file(codeblock, code_text_filename, 'w')

            else:
                while (not code_execution_success):
                    code, code_text_filename, code_filename = get_asset_explanation_gpt4v(table, document_path, gpt4v_prompt = table_code_description_prompt, prompt_extension=prompt_extension, with_context=True, extension='.codeblock', temperature=temperature, model_info=model_info)
                    code_execution_success, exception, output = execute_python_code_block(code_filename)
                    if code_execution_success: 
                        description = f"Python Code executed successfully for table {index} on chunk {chunk_number}\n\nOutput:\n{output}\n"
                        logc(f"Table Post-Processing Success", description)
                        with open(code_filename + '.execution_ok.txt', 'w', encoding='utf-8') as file:
                            file.write(description)
                        break

                    prompt_extension = "\nThe previous code generation failed with the following error:\n\n" + str(exception) + "\n\nPlease fix the error and try again.\n\n"
                    description = f"Extracted Code for table {index} on chunk {chunk_number} could not be executed properly.\n\nCode: {code}\n\nError: {exception}\n\n"
                    logc(f"Table Post-Processing Error. Retry {retries+1}/5", description)
                    temperature += 0.1
                    retries += 1
                    if retries > 4: 
                        description = f"Extracted Code for table {index} on chunk {chunk_number} could not be executed properly.\n\nCode: {code}\n\nError: {exception}\n\n"
                        with open(code_filename + '.execution_errorlog.txt', 'w', encoding='utf-8') as file:
                            file.write(description)
                        break
                
            text = remove_code(text)
            write_to_file(text, text_filename, 'w')
            table_text_files.append(text_filename)
            table_code_text_filenames.append(code_text_filename)
            table_code_py_filenames.append(code_filename)
            table_markdown_filenames.append(markdown_filename)

            # write_to_file(f'\n\n\n#### START OF DESCRIPTION OF TABLE {index}\n' + remove_code(text) + '\n#### END OF DESCRIPTION OF TABLE \n\n', chunk_text_file, mode='a')
            write_to_file(f'\n\n\n#### START OF DESCRIPTION OF TABLE {index}\n' + remove_code(text) + f'\n#### END OF DESCRIPTION OF TABLE {index}\n\n', master_text_file, mode='a')
            # write_to_file(remove_code(text) + '\n\n', master_text_file, mode='a')

            
            time.sleep(2)
            optimized_tag_list = generate_tag_list(remove_code(text), model = model_info['AZURE_OPENAI_MODEL'], client = client)
            write_to_file(optimized_tag_list, replace_extension(text_filename, '.tags.txt'))

        else:
            logc(f"Table Tags File Already Exists for file {table}")
            text_filename = replace_extension(table, '.txt')
            code_filename = text_filename.replace('.txt', '.py')
            if os.path.exists(code_filename): table_code_py_filenames.append(code_filename)
            code_text_filename = text_filename.replace('.txt', '.codeblock')
            if os.path.exists(code_text_filename): table_code_text_filenames.append(code_text_filename)
            markdown_filename = text_filename.replace('.txt', '.md')
            if os.path.exists(markdown_filename): table_markdown_filenames.append(markdown_filename)
            table_text_files.append(text_filename)



    logc(f"Post-Processing: Table processed in chunk {chunk_number} using {model_info['AZURE_OPENAI_RESOURCE']}")
    chunk_dict['table_py'] = table_code_py_filenames
    chunk_dict['table_codeblock'] = table_code_text_filenames
    chunk_dict['table_text_files'] = table_text_files
    chunk_dict['table_markdown'] = table_markdown_filenames
    chunk_dict['text_file'] = text_filename


    return [{'table_py':table_code_py_filenames, 'table_codeblock':table_code_text_filenames, 'table_text':table_text_files, 'table_markdown':table_markdown_filenames}]

def get_ingested_document_text_files(directory, excluded_endings = ['.original.txt', '.tags.txt', '.processed.txt', '.execution_ok.txt', '.execution_errorlog.txt']):
    text_files = []
    image_text_files = []
    table_text_files = []

    # Define the subfolders
    subfolders = ['text', 'images', 'tables']

    # Walk through the directory and its subdirectories
    for subfolder in subfolders:
        for file in glob.glob(os.path.join(directory, subfolder, '*.txt')):
            # Check if the file ends with any of the excluded endings
            if not any(file.endswith(ending) for ending in excluded_endings):
                # Check which subfolder we're in and add the file to the appropriate list
                if subfolder == 'text':
                    text_files.append(file)
                elif subfolder == 'images':
                    image_text_files.append(file)
                elif subfolder == 'tables':
                    table_text_files.append(file)

    return text_files, image_text_files, table_text_files

def get_ingested_document_jpg_images(directory):
    jpg_images = []

    # Define the subfolder
    subfolder = 'images'

    # Walk through the directory and its subdirectory
    for file in glob.glob(os.path.join(directory, subfolder, '*.jpg')):
        jpg_images.append(file)

    return jpg_images

def get_ingested_document_png_table_images(directory):
    png_images = []

    # Define the subfolder
    subfolder = 'tables'

    # Walk through the directory and its subdirectory
    for file in glob.glob(os.path.join(directory, subfolder, '*.png')):
        png_images.append(file)

    return png_images

def create_ingestion_pipeline_dict(ingestion_params_dict): 

    doc_path = ingestion_params_dict['doc_path'] 
    index_name = ingestion_params_dict['index_name']
    
    ingestion_directory = ingestion_params_dict.get('ingestion_directory', None)
    num_threads = ingestion_params_dict.get('num_threads', len([1 for x in gpt4_models if x['AZURE_OPENAI_RESOURCE'] is not None]))
    password = ingestion_params_dict.get('password', None)
    models = ingestion_params_dict.get('models', gpt4_models)
    vision_models = ingestion_params_dict.get('models', gpt4_models)
    delete_existing_output_dir = ingestion_params_dict.get('delete_existing_output_dir', False)
    chunk_size = int(ingestion_params_dict.get('chunk_size', 512))
    chunk_overlap = int(ingestion_params_dict.get('chunk_overlap', 128))

    if ingestion_directory is None:
        ingestion_directory = os.path.join(ROOT_PATH_INGESTION, index_name)

    # Create the Ingestion directory if it doesn't exist
    os.makedirs(ingestion_directory, exist_ok=True)

    download_directory = os.path.join(ingestion_directory, 'downloads')
    os.makedirs(download_directory, exist_ok=True)

    assets = {}
    
    if is_file_or_url(doc_path) == 'url':
        doc_path = download_file(doc_path, os.path.join(ingestion_directory, 'downloads'))
        if doc_path is None:
            return assets

    # Execute file operations based on the file extension
    base_name = os.path.splitext(os.path.basename(doc_path))[0].strip()
    try:
        extension = os.path.splitext(os.path.basename(doc_path))[1].strip()
    except:
        extension = ''


    # Create the directories if it doesn't exist
    doc_proc_directory = os.path.join(ingestion_directory, base_name+extension).replace(" ", "_")

    if delete_existing_output_dir: # and (processing_plan is None): 
        logc("Ingestion Directory remove flag is True. Trying to delete stuff ..")
        shutil.rmtree(doc_proc_directory, ignore_errors=True)
        complete_flag = os.path.join(download_directory, base_name+extension+'.ingested').replace(" ", "_")
        if os.path.exists(complete_flag): 
            os.remove(complete_flag)
            logc(f"Deleting Completed flag from Downloads directory {complete_flag}")

    os.makedirs(doc_proc_directory, exist_ok=True)

    print("Dirname", os.path.dirname(ingestion_directory))
    print("Doc Path: ", doc_path)
    print("Doc Proc Directory: ", doc_proc_directory)
    print("Ingestion Directory: ", ingestion_directory)
    print("Basename: ", base_name)
    print("Extension: ", extension)

    # if extension == '.docx':
    #     document_path = save_docx_as_pdf(doc_path, doc_proc_directory)
    # elif extension == '.pptx':
    #     document_path = save_pptx_as_pdf(doc_path, doc_proc_directory)
    # elif extension == '.pdf':
    #     document_path = os.path.join(doc_proc_directory, base_name + '.pdf').replace(" ", "_")
    #     shutil.copyfile(doc_path, document_path)
    # else:
    #     raise ValueError('Unsupported file extension: {}'.format(extension))

    document_path = os.path.join(doc_proc_directory, os.path.basename(doc_path)).replace(" ", "_")
    shutil.copyfile(doc_path, document_path)

    print("PDF Path: ", document_path)
    master_py_filename = os.path.join(doc_proc_directory, base_name + '.py')

    if extension != '.txt':
        full_text_filename = os.path.join(doc_proc_directory, base_name + '.txt')
    else:
        full_text_filename = os.path.join(doc_proc_directory, base_name + '.txt.txt')

    master_py_filename = master_py_filename.replace(' ', '_')
    full_text_filename = full_text_filename.replace(' ', '_')

    # document_id = str(uuid.uuid4())
    unique_identifier = f"{index_name}_{os.path.basename(doc_path)}"
    document_id = generate_uuid_from_string(unique_identifier)

     # Directory to save text, high-resolution chunk images, and images
    chunks_as_images_directory = os.path.join(os.path.dirname(document_path), 'chunk_images')
    images_directory = os.path.join(os.path.dirname(document_path), 'images')
    text_directory = os.path.join(os.path.dirname(document_path), 'text')
    tables_directory = os.path.join(os.path.dirname(document_path), 'tables')

    # Create the directory if it doesn't exist
    os.makedirs(chunks_as_images_directory, exist_ok=True)
    os.makedirs(images_directory, exist_ok=True)
    os.makedirs(text_directory, exist_ok=True)
    os.makedirs(tables_directory, exist_ok=True)


    ingestion_pipeline_dict = {
        'document_id': document_id,
        'original_document_path': doc_path,
        'original_document_filename': os.path.basename(doc_path),
        'original_document_extension': extension,
        'index_name': index_name,
        'document_processing_directory': doc_proc_directory,
        'document_ingestion_directory': ingestion_directory,
        'document_downloads_directory': download_directory,
        'ingestion_directory': ingestion_directory,
        'download_directory': download_directory,
        'document_path': document_path,
        'master_py_file': master_py_filename,
        'full_text_file': full_text_filename,
        'image_files': [],
        'tables_py': [],
        'tables_md': [],
        'table_files': [],
        'text_files': [],
        'image_text_files': [],
        'table_text_files': [],
        'py_files': [],
        'codeblock_files': [],
        'markdown_files': [],
        'mermaid_files': [],
        "basename": base_name,
        "extension": extension,
        "password": password,
        'extract_text_mode': 'GPT',
        'extract_images_mode': 'GPT',
        'extract_text_from_images': True,
        'models': models,
        'vision_models': vision_models,
        'num_threads': num_threads,
        'chunks_as_images_directory': chunks_as_images_directory,
        'images_directory': images_directory,
        'text_directory': text_directory,
        'tables_directory': tables_directory,
        'chunk_size': int(chunk_size),
        'chunk_overlap': int(chunk_overlap),
    }

    for v, k in ingestion_params_dict.items():
        if v not in ingestion_pipeline_dict:
            ingestion_pipeline_dict[v] = k

    return ingestion_pipeline_dict

def delete_pdf_chunks(ingestion_pipeline_dict):
    # pdf_document chunk
    del ingestion_pipeline_dict['pdf_document']
    for p in ingestion_pipeline_dict['chunks']: del p['chunk']
    return ingestion_pipeline_dict

def process_pdf(ingestion_pipeline_dict, password = None, extract_text_mode = "GPT", extract_images_mode = "GPT", extract_text_from_images = True, models = gpt4_models, vision_models=gpt4_models, num_threads=4, processing_plan=None, verbose=False):
    # print(ingestion_pipeline_dict)
    document_file_path = ingestion_pipeline_dict['document_path']

    # Open the PDF file
    pdf_document = fitz.open(document_file_path)
    full_basename = os.path.basename(document_file_path)

    # Execute file operations based on the file extension
    base_name = os.path.splitext(os.path.basename(document_file_path))[0].strip()
    try:
        extension = os.path.splitext(os.path.basename(document_file_path))[1].strip()
    except:
        extension = ''


    if password is not None: 
        r = pdf_document.authenticate(password)
        if r == 0: raise ValueError("Password is incorrect.")
        filename = document_file_path+'.decrypted.pdf'
        pdf_document.save(filename)
        logc(f"Ingestion Stage of {full_basename}- Info", f"Opened the file with the password. Status: {r}", verbose=verbose)


    logc(f"Ingestion Stage of {full_basename} - Info", f"PDF File with num chunks -> {len(pdf_document)}", verbose=verbose)

    ingestion_pipeline_dict['num_chunks'] = len(pdf_document)
    ingestion_pipeline_dict['document_file_path'] = document_file_path
    ingestion_pipeline_dict['pdf_document'] = pdf_document
    
    ingestion_pipeline_dict['chunks'] = [{
        'chunk':chunk, 
        'chunk_number':index+1, 
        'full_chunk_text':'',
        'images': [],
        'tables': [],
        'image_py': [],
        'image_codeblock': [],
        'image_markdown': [],
        'image_mm': [],
        'image_text': [],
        'table_text': [],
        'table_py': [],
        'table_codeblock': [],
        'table_markdown': [],        
    } for index, chunk in enumerate(pdf_document)]

    doc_proc_directory = ingestion_pipeline_dict['document_processing_directory'] 

    
    if processing_plan is None:
        processing_plan = ['pdf_extract_high_res_chunk_images', 'pdf_extract_text', 'harvest_code', 'pdf_extract_images', 'post_process_images', 'extract_tables_from_images', 'post_process_tables']
    else: 
        text_files, image_text_files, table_text_files = get_ingested_document_text_files(ingestion_pipeline_dict['document_processing_directory'])
        ingestion_pipeline_dict['text_files'] = text_files
        ingestion_pipeline_dict['image_text_files'] = image_text_files
        ingestion_pipeline_dict['table_text_files'] = table_text_files

        image_files = get_ingested_document_jpg_images(ingestion_pipeline_dict['document_processing_directory'])
        table_files = get_ingested_document_png_table_images(ingestion_pipeline_dict['document_processing_directory'])

        for chunk_dict in ingestion_pipeline_dict['chunks']:
            chunk_number = chunk_dict['chunk_number']
            if 'pdf_extract_text' not in processing_plan: chunk_dict['text_file'] = [f for f in text_files if f.endswith(f'chunk_{chunk_number}.txt')][0]
            if 'pdf_extract_images' not in processing_plan: chunk_dict['images'] = [f for f in image_files if (os.path.basename(f).startswith(f'chunk_{chunk_number}_image_')) and (os.path.basename(f).endswith('.jpg'))]
            if 'extract_tables_from_images' not in processing_plan: chunk_dict['tables'] = [f for f in table_files if (os.path.basename(f).startswith(f'chunk_{chunk_number}_table_')) and (os.path.basename(f).endswith('.png'))]


    logc(f"Ingestion Stage 1/7 of {full_basename}", f"Extracting High-Resolution PNG Images from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'pdf_extract_high_res_chunk_images' in processing_plan:
        ingestion_pipeline_dict = pdf_extract_high_res_chunk_images(ingestion_pipeline_dict)

    logc(f"Ingestion Stage 2/7 of {full_basename}", f"Extracting Text with Extract Mode {extract_text_mode}", verbose=verbose)
    if 'pdf_extract_text' in processing_plan:
        ingestion_pipeline_dict = pdf_extract_text(ingestion_pipeline_dict)
    
    logc(f"Ingestion Stage 3/7 of {full_basename}", f"Harvesting Code from Text from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'harvest_code' in processing_plan:
        ingestion_pipeline_dict = harvest_code(ingestion_pipeline_dict)
    
    logc(f"Ingestion Stage 4/7 of {full_basename}", f"Detecting and Extracting Images from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'pdf_extract_images' in processing_plan:
        ingestion_pipeline_dict = pdf_extract_images(ingestion_pipeline_dict)

    ingestion_pipeline_dict = delete_pdf_chunks(ingestion_pipeline_dict)

    logc(f"Ingestion Stage 5/7 of {full_basename}", f"Post-Processing extracted Images from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'post_process_images' in processing_plan:
        ingestion_pipeline_dict = post_process_images(ingestion_pipeline_dict)

    logc(f"Ingestion Stage 6/7 of {full_basename}", f"Detecting and Extracting Tables from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'extract_tables_from_images' in processing_plan:
        ingestion_pipeline_dict = extract_tables_from_images(ingestion_pipeline_dict)

    logc(f"Ingestion Stage 7/7 of {full_basename}", f"Post-Processing extracted Tables from PDF with {len(pdf_document)} chunks", verbose=verbose)
    if 'post_process_tables' in processing_plan:
        ingestion_pipeline_dict = post_process_tables(ingestion_pipeline_dict)

    # Close the PDF document
    pdf_document.close()

    vec_entries = len(ingestion_pipeline_dict['text_files'] + ingestion_pipeline_dict['image_text_files'] + ingestion_pipeline_dict['table_text_files'])

    pkl_path = os.path.join(doc_proc_directory, f'{base_name}.pkl')

    if os.path.exists(pkl_path):
        ext = f".{get_current_time()}.pkl"
        shutil.copyfile(pkl_path, pkl_path + ext)

    save_to_pickle(ingestion_pipeline_dict, pkl_path)


    logc(f"Ingestion Stage of {base_name} Complete", f"{get_current_time()}: Ingestion of document {base_name} resulted in {vec_entries} entries in the Vector Store", verbose=verbose)
    return ingestion_pipeline_dict















def create_metadata(asset_file, file_id, document_path, document_id, asset_type="text", image_file = "", python_block = "", python_code = "", markdown = "", mermaid_code = "", tags = ""):
    metadata = {
        "asset_path": asset_file, 
        "document_path": document_path, 
        "filename": os.path.basename(document_path),
        "image_file": image_file,
        "asset_filename": asset_file,
        "chunk_number": extract_chunk_number(asset_file),
        "type": asset_type,
        "document_id": document_id,
        "python_block" : python_block,
        "python_code" : python_code,
        "markdown": markdown,
        "mermaid": mermaid_code,
        "tags": tags,
        "asset_id": file_id
    }

    for m in metadata:
        metadata[m] = metadata[m].replace("\\", "/")

    return metadata



def generate_tag_list(text, model = AZURE_OPENAI_MODEL, client = oai_client):
    try:
        messages = [{"role":"system", "content":optimize_embeddings_prompt.format(text=text)}]
        result = get_chat_completion(messages, model=model, client = client) 
        return result.choices[0].message.content
    except Exception as e:
        logc("Error generating tag list: ", e)
        return text



def generate_document_wide_tags(ingestion_pipeline_dict):
    doc_proc_directory = ingestion_pipeline_dict['document_processing_directory']
    basename = ingestion_pipeline_dict['basename']
    full_text_file = ingestion_pipeline_dict['full_text_file']
    full_text = read_asset_file(full_text_file)[0]

    tags_text_file = os.path.join(doc_proc_directory, f"{basename}.tags.txt")
    text = limit_token_count(full_text)

    prompt = document_wide_tags.format(text=text)
    tags = ask_LLM(prompt)

    write_to_file(tags, tags_text_file, 'w')
    ingestion_pipeline_dict['document_wide_tags'] = tags

    return ingestion_pipeline_dict



def generate_document_wide_summary(ingestion_pipeline_dict):
    doc_proc_directory = ingestion_pipeline_dict['document_processing_directory']
    basename = ingestion_pipeline_dict['basename']
    full_text_file = ingestion_pipeline_dict['full_text_file']
    full_text = read_asset_file(full_text_file)[0]

    summary_text_file = os.path.join(doc_proc_directory, f"{basename}.summary.txt")
    text = limit_token_count(full_text)

    prompt = document_wide_summary.format(text=text)
    summary = ask_LLM(prompt)

    write_to_file(summary, summary_text_file, 'w')
    ingestion_pipeline_dict['document_wide_summary'] = summary

    return ingestion_pipeline_dict

def add_asset_to_vec_store(assets, index, asset_file, document_path, document_id, vector_type = "AISearch"):

    text = ""
    python_code = ""
    python_block = ""
    markdown = "" 
    image_file = ""
    mermaid_code = ""
    tags = ""
    doc_proc_directory = assets['document_processing_directory']
    basename = assets['basename']
    original_document_filename = assets['original_document_filename']
    index_name = assets['index_name']

    tags_text_file = os.path.join(doc_proc_directory, f"{basename}.tags.txt")

    # asset_file = os.path.abspath(asset_file)
    # document_path = os.path.abspath(document_path)

    if "image" in asset_file:
        asset_type = "image"
        text, status = read_asset_file(asset_file)

        image_file = check_replace_extension(asset_file, '.jpg')
        python_code = check_replace_extension(asset_file, '.py')
        mermaid_code = check_replace_extension(asset_file, '.mermaid')
        python_block = check_replace_extension(asset_file, '.codeblock')

    elif "table" in asset_file:
        asset_type = "table"
        text, status = read_asset_file(replace_extension(asset_file, '.txt'))

        python_block = check_replace_extension(asset_file, '.codeblock')
        python_code = check_replace_extension(asset_file, '.py')
        markdown = check_replace_extension(asset_file, '.md')
        image_file = check_replace_extension(asset_file, '.png')

    else:
        asset_type = "text"
        text, status = read_asset_file(asset_file)
        
        python_block = check_replace_extension(asset_file, '.codeblock')
        python_code = check_replace_extension(asset_file, '.py')
        markdown = check_replace_extension(asset_file, '.md')
        
    tags = ''

    if os.path.exists(tags_text_file):
        logc("Document-Wide Tags Added")
        tags += read_asset_file(tags_text_file)[0]

    tags_file = check_replace_extension(asset_file, '.tags.txt')
    logc(f"local tags file {(tags_file != '') and (os.path.exists(tags_file))}", tags_file)
    if (tags_file != "") and (os.path.exists(tags_file)):
        logc("Chunk-specific Tags Added")
        tags += ', ' + read_asset_file(tags_file)[0]


    # file_id = str(uuid.uuid4())
    unique_identifier = f"{index_name}_{original_document_filename}_{os.path.basename(asset_file)}"
    file_id = generate_uuid_from_string(unique_identifier)

    metadata = create_metadata(asset_file, file_id, document_path, document_id, asset_type=asset_type, image_file = image_file, python_block = python_block, python_code = python_code, markdown = markdown, mermaid_code=mermaid_code, tags=tags)
    print(f"\n{bc.OKBLUE}Metadata:\n{bc.OKGREEN}{json.dumps(metadata, indent=4)}\n{bc.ENDC}")
    

    if asset_type == "text":
        chunk_number = extract_chunk_number(asset_file)
        text_for_embeddings = get_processed_context_chunks(asset_file, text, int(chunk_number))
    else: 
        chunk_number = extract_chunk_number(asset_file)
        text_for_embeddings = get_processed_context_chunk(doc_proc_directory, text, int(chunk_number))

    if vector_type == "AISearch":
        metadata['text'] = text
        metadata['vector'] = get_embeddings(text_for_embeddings)
        index.upload_documents([metadata])

    return file_id, metadata

def check_vecstore_vs_ingestion_dir(index_name):
    files = glob.glob(f'{index_name}/**/*.vecstore.txt', recursive=True)
    cogrequest = CogSearchRestAPI(index_name = index_name)
    doc_count_in_vecstore = cogrequest.get_stats()['documentCount']

    doc_count_in_dir = 0

    for f in files:
        items = json.loads(read_asset_file(f)[0])
        doc_count_in_dir += len(items)

    print(f"In VecStore: {doc_count_in_vecstore}. In Directory: {doc_count_in_dir}. Difference: {doc_count_in_vecstore - doc_count_in_dir}")
    return doc_count_in_vecstore, doc_count_in_dir, (doc_count_in_vecstore - doc_count_in_dir)

def commit_assets_to_vector_index(assets, ingestion_directory, index_name = 'mm_doc_analysis', vector_type = "AISearch"):

    text_files = assets['text_files']
    image_text_files = assets['image_text_files']
    table_text_files = assets['table_text_files']
    document_path = assets['document_path']
    # document_path = assets['original_document_path']
    document_id = assets['document_id']
    # logc("Assets: ", assets, verbose=True)


    if vector_type == "AISearch":
        index = CogSearchRestAPI(index_name)
        if index.get_index() is None:
            print(f"No index {index_name} detected, creating one ... ")
            index.create_index()

    index_ids = []
    metadatas = []
    for asset_file in text_files + image_text_files + table_text_files:
        asset_file_id, metadata = add_asset_to_vec_store(assets, index, asset_file, document_path, document_id, vector_type=vector_type)
        print("Asset File ID: ", asset_file_id)
        index_ids.append(asset_file_id)
        metadatas.append(metadata)

    return index_ids, metadatas

def get_processed_context_chunks(asset_file, text, chunk_number):
    
    dir_name = os.path.dirname(asset_file)

    try:
        previous_chunk_text = read_asset_file(os.path.join(dir_name, f"chunk_{chunk_number-1}.processed.txt"))[0]
    except:
        previous_chunk_text = ""

    try:
        next_chunk_text = read_asset_file(os.path.join(dir_name, f"chunk_{chunk_number+1}.processed.txt"))[0]
    except:
        next_chunk_text = ""

    return previous_chunk_text + "\n\n" + text + "\n\n" + next_chunk_text

def get_processed_context_chunk(doc_proc_directory, text, chunk_number):
    
    path = os.path.join(doc_proc_directory, f"text/chunk_{chunk_number}.processed.txt")
    try:
        current_chunk_text = read_asset_file(path)[0]
    except:
        current_chunk_text = ""

    return text + "\n\n" + current_chunk_text

def get_context_chunks(document_path, chunk_number):
    try:
        pdf_document = fitz.open(document_path)
        try:
            if chunk_number-2 >= 0:
                previous_chunk = pdf_document[chunk_number-2].get_text()
            else:
                previous_chunk = 0
        except:
            previous_chunk = ""

        try:
            current_chunk = pdf_document[chunk_number-1].get_text()
        except:
            current_chunk = ""
        
        try:
            next_chunk = pdf_document[chunk_number].get_text()
        except:
            next_chunk = ""

        pdf_document.close()

        return previous_chunk, current_chunk, next_chunk
    except:
        return "No chunk retrieved.", "No chunk retrieved.", "No chunk retrieved."

def get_asset_explanation_gpt4v(asset_file, document_path, gpt4v_prompt = image_description_prompt, prompt_extension = "", with_context = False, extension = None, temperature = 0.2, model_info=None):

    code_filename = ''
    text_filename = ''
    prompt_ext = ''

    if with_context:
        chunk_number = extract_chunk_number(asset_file)
        previous_chunk, current_chunk, next_chunk = get_context_chunks(document_path, int(chunk_number))
        prompt_ext = prompt_extension + context_extension.format(previous_chunk = previous_chunk, current_chunk = current_chunk, next_chunk = next_chunk)
    
    try:
        text, description = call_gpt4v(asset_file, gpt4v_prompt = gpt4v_prompt, prompt_extension = prompt_ext, temperature = temperature, model_info=model_info)
    except Exception as e:
        logc(f"get_asset_explanation_gpt4v:: Error generating text for asset: {asset_file}\nError: {e}")
        text = "No results could be extracted or explanation generated due to API errors."
        description = f"Error generating text for asset: {asset_file}\nError: {e}"
    

    if extension == 'dont_save':
        pass
    elif extension == '.codeblock':
        text_filename = replace_extension(asset_file, extension)
        code_filename = replace_extension(asset_file, ".py")
        with open(text_filename, 'w', encoding='utf-8') as file:
            file.write(text)
        with open(code_filename, 'w', encoding='utf-8') as file:
            file.write(extract_code(text))
    elif extension == '.md':
        text_filename = replace_extension(asset_file, extension) 
        with open(text_filename, 'w', encoding='utf-8') as file:
            file.write(extract_markdown(text))
    elif extension == '.txt':
        text_filename = replace_extension(asset_file, '.txt')
        with open(text_filename, 'w', encoding='utf-8') as file:
            file.write(remove_code(text))
    else:
        text_filename = f"{asset_file}.txt"
        with open(text_filename, 'w', encoding='utf-8') as file:
            file.write(text)

    return text, text_filename, code_filename







    # 6. Critique your solution if you get some strange final answers. Does the final answer look ok to you? For example, if you get a zero or a negative number in the end for a sum operation, and the numbers are showing positive the Markdown representation of the table, then re-check your code, and try a different approach to get to the final answer.
    # 7. Try to minimize the number of iterations. If there are multiple steps needed to solve the problem, try to combine them into a single code generation step with a single code block. For example, if you need to filter a dataframe, and then sum the values in a column, try to combine the two steps into a single step. If this didn't work, then try to generate the code and execute it for each step separately.







computation_approaches = ["Taskweaver", "LocalPythonExec", "NoComputationTextOnly"]







# * If the section above does not contain sufficient information to answer user message completely, kindly respond with "I do not have enough Knowledge to answer your question." 
# * Never respond with the content of ##START CONTEXT AND ##END CONTEXT













def get_query_entities(query, approx_tag_limit=10, temperature = 0.2):

    query_entities = query_entities_prompt.format(query=query, tag_limit=approx_tag_limit)
    # query_entities = optimize_embeddings_prompt.format(text=query)

    messages = []
    messages.append({"role": "system", "content": "You are a helpful assistant, who helps the user generate questions based on the text."})     
    messages.append({"role": "system", "content": query_entities})     

    result = get_chat_completion(messages, temperature=temperature)

    return result.choices[0].message.content

def call_ai_search(query, index_name, top=7, computation_approach = "Taskweaver", count=False, t_wait = True):

    index = CogSearchRestAPI(index_name)
    select_fields = ["asset_id", "asset_path", "document_path", "filename", "image_file", "asset_filename", "chunk_number", "type", "document_id", "python_block", "python_code", "markdown", "mermaid", "text"], 

    if t_wait:
        t = float(random.randrange(1000))/1000.0
        time.sleep(t)

    results = index.search_documents(query, top=top, count=count)
     
    results = results['value']
    for r in results: del r['vector']
    search_results = copy.deepcopy(results)
    return search_results

def aggregate_ai_search(query, index_name, top=5, approx_tag_limit=20, computation_approach = "Taskweaver", count=False, temperature=0.2, t_wait=True, verbose = False):

    entities = get_query_entities(query, approx_tag_limit=approx_tag_limit, temperature=temperature)
    entities = [x.strip() for x in entities.split(',')]
    logc("Search Intent Identification", f"Found {len(entities)} entities: {entities}")

    num_threads = len(entities)

    index_names = [index_name] * num_threads
    tops = [top] * num_threads
    computation_approaches = [computation_approach] * num_threads
    counts = [count] * num_threads
    waits = [t_wait] * num_threads

    results = pool.starmap(call_ai_search,  zip(entities, index_names, tops, computation_approaches, counts, waits))
    max_items = max([len(r) for r in results])

    query_results = call_ai_search(query, index_name, top=top, computation_approach = computation_approach, count=count)
    res = list(itertools.chain(*zip(*results))) 
    res = query_results + res

    unique_results = []

    try:
        for result in res:
            if result['asset_path'] not in [r['asset_path'] for r in unique_results]:
                unique_results.append(result)
    except:
        for result in res:
            if result['text'] not in [r['text'] for r in unique_results]:
                unique_results.append(result)

    if verbose: logc("Found the following asset results:", [r['asset_path'] for r in unique_results])
    # for r in unique_results: logc(r['asset_path'])

    return unique_results

def check_if_computation_is_needed(query):
    messages = []
    messages.append({"role": "system", "content": "You are a helpful AI assistant. You help users answer their queries based on the information supplied below."})     
    messages.append({"role": "user", "content": computation_is_needed_prompt.format(query=query)})   

    result = get_chat_completion(messages)
    return result.choices[0].message.content
     
def apply_computation_support(query, assets, computation_approach="AssistantsAPI", conversation_history = [], user_id = None, include_master_py=True, verbose = False):
    files = []
    if computation_approach == "Taskweaver":
        computation_support, files = try_code_interpreter_for_tables_using_taskweaver(assets, query, include_master_py=include_master_py,verbose = verbose)
    elif computation_approach == "LocalPythonExec":
        computation_support, files = try_code_interpreter_for_tables_using_python_exec(assets, query, include_master_py=include_master_py, verbose = verbose)
    elif computation_approach == "AssistantsAPI":
        computation_support, files = try_code_interpreter_for_tables_using_assistants_api(assets, query, user_id = user_id, include_master_py=include_master_py, verbose = verbose)
    else:
        computation_support = "No computation results."

    return computation_support, files



def generate_search_assets(all_results, limit = 1000, verbose=False):
    assets = {}
    assets['python_block'] = []
    assets['python_code'] = []
    assets['filenames'] = []
    assets['asset_filenames'] = []
    assets['document_paths'] = []
    assets['vision_images'] = []

    # logc("All Results", all_results)
    results = all_results[:limit]

    # logc("Metadatas", json.dumps(results, indent=4), verbose = verbose)
    if verbose: logc("Search Function Executing...", f"Found {len(results)} search results")


    for metadata in results:
        if metadata['type'] == 'table':
            assets['filenames'].append(metadata['filename'])
            assets['asset_filenames'].append(metadata['asset_filename'])
            assets['python_block'].append(metadata['python_block'])
            assets['python_code'].append(metadata['python_code'])
            assets['document_paths'].append(metadata['document_path'])

        elif (metadata['type'] == 'image'):
            assets['filenames'].append(metadata['filename'])
            assets['asset_filenames'].append(metadata['asset_filename'])
            if metadata['python_block'] == "":
                assets['python_block'].append(metadata['asset_filename'])
            else:
                assets['python_block'].append(metadata['python_block'])
            assets['python_code'].append(metadata['python_code'])
            assets['document_paths'].append(metadata['document_path'])
            assets['vision_images'].append({'pdf':metadata['document_path'], 'img':metadata['image_file']})


        elif (metadata['type'] == 'text') and (metadata['python_block'] != ""):
            assets['filenames'].append(metadata['filename'])
            assets['asset_filenames'].append(metadata['asset_filename'])
            assets['python_block'].append(metadata['python_block'])
            assets['python_code'].append(metadata['python_code'])
            assets['document_paths'].append(metadata['document_path'])
    
    return assets



def detect_intent_of_query(query):
    prompt = detect_intent_prompt.format(query=query, history=get_history_as_string([]))
    answer = ask_LLM_with_JSON(prompt)
    answer_dict = recover_json(answer)
    print(answer_dict)
    return answer_dict['category']
    
def get_history_as_string(conversation_history):
    history = ''

    for c in conversation_history:
        content = c['content'].replace('\n\n', ' ')
        history += f"## {c['role']}: {content}\n\n"

    return history

def search(query, learnings = None, top=7, approx_tag_limit=15, conversation_history = [], user_id = None, computation_approach = "AssistantsAPI", computation_decision = "LLM", vision_support = False, include_master_py=True, vector_directory = None, vector_type = "AISearch", index_name = 'mm_doc_analysis', full_search_output = True, count=False, token_limit = 100000, temperature = 0.2, verbose = False):
    global search_context_extension, search_system_prompt, search_prompt #FIXME: Remove this global variable

    if vector_directory is None:
        vector_directory = os.path.join(ROOT_PATH_INGESTION, index_name)

    if verbose: logc("Search Function Executing...", "Starting Search ...")   

    vision_support_result = "No vision results"
    computation_support = "No computation results."

    search_results = {}
    files = []

    results = []
    if vector_type == "AISearch":
        results = aggregate_ai_search(query, index_name, top=top, approx_tag_limit=approx_tag_limit, computation_approach=computation_approach, count=count, temperature=temperature, verbose = verbose)
        text_results = [result['text'] for result in results] # FIXME


    ## Limit the results
    # results = results[:35]
    assets = generate_search_assets(results, verbose = verbose)

    summaries = []

    for doc in assets['document_paths']:
        summary = read_asset_file(replace_extension(doc, ".summary.txt"))[0] 
        if summary != '':
            summary_prompt = summaries_context_extension.format(proc_filename=os.path.basename(doc), summary=summary)
            summaries.append(summary_prompt)

    summaries = list(set(summaries))

    logc("Search Results", {"results":results}, verbose = verbose)

    if vision_support:
        vision_support_result = ""

        img_counter = 0
        for p in assets['vision_images']:
            document_path = p['pdf']
            img_path = p['img']

            try:
                interm_vision_support_result, _, _ = get_asset_explanation_gpt4v(img_path, document_path, gpt4v_prompt = vision_support_prompt.format(query=query), with_context = True, extension = "dont_save")

                vision_support_result += f"## START OF VISION RESULT\nPDF: {os.path.basename(document_path)}\nImage: {os.path.basename(img_path)}\nAnswer from Image:\n{interm_vision_support_result}\n## END OF VISION RESULT\n\n"
                img_counter += 1
            except Exception as e:
                print(f"Error processing vision support: {e}")

        if vision_support_result == "": vision_support_result = "No vision results."
        if verbose: logc("Vision Support Results", vision_support_result, verbose=verbose)


    logc("Assets", assets)
    


    if computation_approach != "NoComputationTextOnly":
        if computation_decision == "LLM":
            # logc("Checking Computation Intent", verbose = verbose)
            intent = check_if_computation_is_needed(query)
            logc("Computation Intent", intent)

            if "YES" in intent:
                computation_support, files = apply_computation_support(query, assets, computation_approach, conversation_history = conversation_history, user_id = user_id, include_master_py=include_master_py, verbose = verbose)
                logc("Computation Support Output", computation_support)
                
            
        elif computation_decision == "Force":
            computation_support, files = apply_computation_support(query, assets, computation_approach, conversation_history = conversation_history, user_id = user_id,include_master_py=include_master_py, verbose = verbose)
            if verbose: logc("Search Function Executing...", f"Computation Support Output\n{computation_support}")


    unique_results = []

    for result in results:
        if result['asset_path'] not in [r['asset_path'] for r in unique_results]:
            unique_results.append(result)


    # context_array = [search_context_extension.format(   
    #         number = index,
    #         search_result = clean_up_text(result['text']), 
    #         filename = os.path.relpath(result['asset_path']),
    #         proc_filename = os.path.basename(result['document_path']),
    #         document_path = os.path.relpath(result['document_path']),
    #         type = result['type'],
    #         chunk_number = result['chunk_number']) for index, result in enumerate(unique_results)]


    context_array = []

    for index, result in enumerate(unique_results):
        analysis = read_asset_file(replace_extension(result['asset_path'], ".analysis.txt"))[0]  
        analysis = f"Analysis:\nHere's analysis of this text chunk in relation to the whole document:\n{analysis}" if analysis != "" else ""

        search_context = search_context_extension.format(   
            number = index,
            search_result = clean_up_text(result['text']), 
            filename = os.path.relpath(result['asset_path']),
            proc_filename = os.path.basename(result['document_path']),
            document_path = os.path.relpath(result['document_path']),
            type = result['type'],
            chunk_number = result['chunk_number'],
            analysis = analysis,
        )

        context_array.append(search_context)


    context_window = []
    token_window = 0 

    for e in context_array:
        token_window += get_token_count(e)
        if token_window < token_limit:
            context_window.append(e)
        else:
            break

    context = '\n'.join(context_window)

    document_summaries = ''

    if len(summaries) > 0:
        summaries = '\n'.join(summaries)
        document_summaries = f"**Document Summaries**\n## START OF DOCUMENTSUMMARIES\n\n{summaries}\n\n## END OF DOCUMENT SUMMARIES\n\n"

    logc("Document Summaries", document_summaries)

    logc("Full Context", context)


    if learnings is not None:
        query = search_learnings_template.format(user_query=query, learnings=learnings)
        if verbose: logc("Improved Query", query)
         
    if full_search_output:
        full_search_prompt = search_prompt.format(context=context, query=query, vision_support =  vision_support_result, computation_support=computation_support, search_json_output=full_search_json_output, document_summaries=document_summaries)
    else:
        full_search_prompt = search_prompt.format(context=context, query=query, vision_support =  vision_support_result, computation_support=computation_support, search_json_output=limited_search_json_output, document_summaries=document_summaries)

    logc("Full Search Prompt...", full_search_prompt)

    messages = []
    messages.append({"role": "system", "content": search_system_prompt})     
    messages.extend(conversation_history)
    messages.append({"role": "user", "content": full_search_prompt})     

    logc("Search Function Executing...", f"Seach Query Token Count => {get_token_count(full_search_prompt)}")
    result = get_chat_completion_with_json(messages, temperature=temperature)

    if verbose: logc("Final Prompt", f"{result.choices[0].message.content}")

    final_json = recover_json(result.choices[0].message.content)
    # logc("Final Answer in JSON", final_json)


    for r in final_json['references']:
        num = int(r['search_result_number'])
        result = unique_results[num]

        r['asset'] = result['asset_path']
        r['document_path'] = result['document_path']
        r['original_document'] = result['filename']
        r['section'] = result['chunk_number']
        r['type'] = result['type']
        r['search_result_number'] = num
        r['sas_link'] = generate_file_sas_full_link(result['document_path'])

    logc("Final Answer in JSON", final_json)


    try:
        final_answer = final_json['final_answer']
    except:
        final_answer = "No final answer."

    try:
        references = final_json['references']
        output_excel = final_json['output_excel_file']
    except:
        references = []
        output_excel = ""

    # conversation_history.append({"role": "user", "content": query})
    # conversation_history.append({"role": "assistant", "content": final_answer})

    # conversation_history= conversation_history[:-6]

    print("Final Answer", final_answer)

    return final_answer, references, output_excel, search_results, files